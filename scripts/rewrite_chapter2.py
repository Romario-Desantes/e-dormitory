from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.build_diploma_report import (  # noqa: E402
    add_canvas_title,
    create_canvas,
    draw_arrow,
    draw_box,
    load_font,
    save_image,
)

DOCX_PATH = ROOT / "labs_reports" / "Diploma_Report_Student_Dormitory.docx"
OUTPUT_COPY = ROOT / "output" / "doc" / "Diploma_Report_Student_Dormitory.docx"


def find_paragraph(document: Document, prefix: str, occurrence: int = 1):
    count = 0
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            count += 1
            if count == occurrence:
                return paragraph
    raise ValueError(f"Paragraph starting with {prefix!r} not found")


def insert_before(paragraph, text: str | None = None):
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_element
    if text is not None:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_after(paragraph, text: str | None = None):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_element
    if text is not None:
        new_paragraph.add_run(text)
    return new_paragraph


def style_heading1(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.0
    if paragraph.runs:
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def style_heading2(paragraph):
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    if paragraph.runs:
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def style_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def style_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    if paragraph.runs:
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_picture_before(paragraph, image_path: Path, width_cm: float = 15.5):
    picture_paragraph = insert_before(paragraph)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return picture_paragraph


def append_block(anchor, texts: list[str]):
    last = anchor
    for text in texts:
        last = insert_after(last, text)
        style_body(last)
    return last


def create_site_map_figure() -> Path:
    image, draw = create_canvas(2200, 1500)
    add_canvas_title(draw, "Структура сайту / застосунку", "Ієрархія сторінок та рольові маршрути")

    title_font = load_font(20, bold=True)
    small_font = load_font(16)
    tiny_font = load_font(14)

    login_box = (810, 140, 1390, 260)
    shell_box = (760, 320, 1440, 470)
    draw_box(draw, login_box, "Login page\n/login\nПеревірка сесії", title_font, fill=(239, 246, 255))
    draw_box(
        draw,
        shell_box,
        "AppShell / /app\nheader + sidebar + notifications\nСпільна оболонка для всіх ролей",
        title_font,
        fill=(236, 253, 245),
        outline=(16, 185, 129),
    )
    draw_arrow(draw, (1100, 260), (1100, 320))

    role_boxes = [
        (120, 600, 650, 790, "Student\n/overview\n/tickets\n/passes\n/finance"),
        (750, 600, 1280, 790, "Commandant\n/occupancy\n/relocations\n/discipline"),
        (1380, 600, 1910, 790, "Master\n/tasks\n/tasks/:ticketId"),
        (250, 940, 780, 1130, "Guard\n/guard"),
        (860, 940, 1390, 1130, "Accountant\n/payments"),
        (1470, 940, 2000, 1130, "Admin\n/students\n/staff\n/directories"),
    ]
    for box in role_boxes:
        draw_box(draw, box[:4], box[4], small_font, fill=(248, 250, 252))

    draw.text((140, 835), "Студент", font=tiny_font, fill=(75, 85, 99))
    draw.text((815, 835), "Комендант", font=tiny_font, fill=(75, 85, 99))
    draw.text((1490, 835), "Майстер", font=tiny_font, fill=(75, 85, 99))
    draw.text((300, 1175), "Охоронець", font=tiny_font, fill=(75, 85, 99))
    draw.text((930, 1175), "Бухгалтер", font=tiny_font, fill=(75, 85, 99))
    draw.text((1550, 1175), "Адміністратор", font=tiny_font, fill=(75, 85, 99))

    arrow_targets = [
        (385, 600, 1090, 470),
        (1015, 600, 1090, 470),
        (1645, 600, 1110, 470),
        (515, 940, 1070, 470),
        (1125, 940, 1100, 470),
        (1735, 940, 1130, 470),
    ]
    for start_x, start_y, end_x, end_y in arrow_targets:
        draw_arrow(draw, (start_x, start_y), (end_x, end_y))

    note_font = load_font(16)
    note_box = (280, 1280, 1920, 1385)
    draw_box(
        draw,
        note_box,
        "Навігація змінюється залежно від ролі, але користувач завжди працює в межах одного каркаса застосунку.\nЦе спрощує орієнтацію, зберігає єдиний стиль і дозволяє легко додавати нові маршрути.",
        note_font,
        fill=(254, 243, 199),
        outline=(245, 158, 11),
    )

    return save_image(image, "figure_2_9_site_map.png")


def update_toc(doc: Document) -> None:
    toc_map = {
        41: "ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ, СКОРОЧЕНЬ І ТЕРМІНІВ      4",
        42: "ВСТУП      5",
        43: "РОЗДІЛ 1 ПОСТАНОВКА ЗАДАЧІ      7",
        44: "1.1 Аналіз предметної області (опис бізнес-процесу)      7",
        45: "1.2 Аналіз існуючих рішень      10",
        46: "1.3 Формування вимог до проєкту      12",
        47: "РОЗДІЛ 2 ПРОЄКТУВАННЯ СИСТЕМИ      15",
        48: "2.1 Проєктування функціоналу та структури системи керування гуртожитком      15",
        49: "2.2 Проєктування бази даних      20",
        50: "2.3 Макети користувацького інтерфейсу      24",
        51: "2.4 Структура сайту/застосунку      28",
        52: "РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ      31",
        53: "3.1 Обґрунтування вибору стеку розробки      31",
        54: "3.2 Реалізація бази даних      33",
        55: "3.3 Реалізація основної функціональності      35",
        56: "3.4 Дизайн користувацького інтерфейсу      37",
        57: "3.5 Розроблення заходів для захисту інформації      39",
        58: "РОЗДІЛ 4 ТЕСТУВАННЯ ТА ВПРОВАДЖЕННЯ      41",
        59: "4.1 Обґрунтування вибору методів тестування      41",
        60: "4.2 Формування тест-плану      42",
        61: "4.3 Тест-кейси та чек-листи      44",
        62: "4.4 Впровадження системи      46",
        63: "ВИСНОВКИ      48",
        64: "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ      49",
    }

    for index, text in toc_map.items():
        doc.paragraphs[index].text = text

    append_block(doc.paragraphs[64], ["ДОДАТКИ      51"])


def main() -> None:
    class_path = ROOT / "output" / "doc" / "figures" / "figure_2_2_class.png"
    sequence_path = ROOT / "output" / "doc" / "figures" / "figure_2_2_sequence.png"
    wireframe_path = ROOT / "output" / "doc" / "figures" / "figure_2_4_wireframe.png"
    site_map_path = create_site_map_figure()

    doc = Document(DOCX_PATH)
    update_toc(doc)

    # Chapter 2 headings.
    find_paragraph(doc, "2.1 Діаграми IDEF0, Data Flow, UML-моделі", 1).text = (
        "2.1 Проєктування функціоналу та структури системи керування гуртожитком"
    )
    find_paragraph(doc, "2.2 ER-модель бази даних", 1).text = "2.2 Проєктування бази даних"
    find_paragraph(doc, "2.3 Макет користувацького інтерфейсу", 1).text = "2.3 Макети користувацького інтерфейсу"
    find_paragraph(doc, "2.4 Структура сайту", 1).text = "2.4 Структура сайту/застосунку"
    for prefix in [
        "2.1 Проєктування",
        "2.2 Проєктування",
        "2.3 Макети",
        "2.4 Структура",
    ]:
        style_heading2(find_paragraph(doc, prefix, 1))

    # Section 2.4 - add site map figure and enrich text first, then go upward.
    sec24 = find_paragraph(doc, "2.4 Структура сайту/застосунку", 1)
    sec24_text1 = insert_after(
        sec24,
        "Структура застосунку побудована за маршрутовим принципом і починається зі сторінки входу. Після успішної авторизації користувач отримує доступ до спільної оболонки, у якій однаково працюють шапка сторінки, ліве меню, панель повідомлень і робоча область. Це дозволяє не дублювати елементи інтерфейсу на кожній сторінці, а зосередити увагу саме на функціональному контенті."
    )
    style_body(sec24_text1)
    sec24_text2 = insert_after(
        sec24_text1,
        "На рівні маршрутизації система використовує окремі гілки для кожної ролі. Студент працює зі сторінками огляду, заявок, перепусток і фінансів; комендант - з кімнатами, переселенням і дисципліною; майстер - з ремонтними заявками; охоронець - з терміналом перевірки перепусток; бухгалтер - з оплатами; адміністратор - з користувачами, персоналом і довідниками. Така декомпозиція спрощує перевірку доступу і робить структуру зрозумілою навіть без додаткових пояснень."
    )
    style_body(sec24_text2)
    sec24_text3 = insert_after(
        sec24_text2,
        "З точки зору фронтенд-архітектури головна сторінка після входу не є окремим набором випадкових екранів. Вона є узгодженим маршрутом, де кожен розділ відкривається за допомогою вкладеної маршрутизації, а активний пункт меню автоматично підсвічується. Для невеликих екранів використовується мобільне меню та нижня навігаційна панель, тому застосунок залишається придатним для роботи як на комп'ютері, так і на смартфоні."
    )
    style_body(sec24_text3)
    sec24_text4 = insert_after(
        sec24_text3,
        "Схема структури сайту показує, що всі сторінки об'єднані не просто за тематикою, а й за поведінкою: єдина сесія, єдина панель сповіщень, загальний хедер і рольовий набір маршрутів формують послідовний користувацький досвід. Такий підхід є зручним для супроводу, оскільки додавання нової функції вимагає лише розширення окремого маршруту і відповідного елемента меню."
    )
    style_body(sec24_text4)
    add_picture_before(find_paragraph(doc, "РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ", 1), site_map_path)
    site_map_caption = insert_after(
        find_paragraph(doc, "РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ", 1)._p.getprevious().getparent().paragraphs[-1]
        if False
        else find_paragraph(doc, "РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ", 1),
        "Рисунок 2.9 - Схема маршрутизації та навігації застосунку",
    )
    # The caption above is inserted directly after the site map figure because the figure paragraph
    # sits immediately before the section 3 heading after add_picture_before().
    style_caption(site_map_caption)

    # Section 2.3 - wireframe and interface layout.
    sec23 = find_paragraph(doc, "2.3 Макети користувацького інтерфейсу", 1)
    sec23_p1 = insert_after(
        sec23,
        "Макети інтерфейсу розробляються не як остаточний дизайн, а як wireframe. Такий підхід дає змогу зосередитися на розміщенні блоків, взаємозв'язках елементів і пріоритетах користувача, не відволікаючись на кольорову схему чи декоративні деталі. Для дипломної роботи це особливо важливо, тому що проєктований сервіс має обслуговувати кілька ролей одночасно і в кожної з них власна логіка роботи."
    )
    style_body(sec23_p1)
    sec23_p2 = insert_after(
        sec23_p1,
        "Базова композиція інтерфейсу складається з хедера, робочої області, навігаційного меню та інформаційних панелей. У верхній частині розміщується назва поточного розділу, коротке позначення ролі, панель сповіщень, дані користувача та кнопка виходу. Ліворуч у десктопній версії передбачено вертикальне меню, а на мобільних пристроях - висувну панель і нижню навігацію, щоб користувач не втрачав доступ до основних дій."
    )
    style_body(sec23_p2)
    sec23_p3 = insert_after(
        sec23_p2,
        "Для студента wireframe включає оглядовий екран із блоком інформації про кімнату, стан фінансового балансу, швидкі дії для подачі заявки та перегляд власних перепусток. Для коменданта окремо спроєктовано екран заселеності, картки кімнат і область перегляду заявок на переселення. Для майстра головним елементом є список ремонтних звернень, а для охоронця - форма перевірки коду перепустки."
    )
    style_body(sec23_p3)
    sec23_p4 = insert_after(
        sec23_p3,
        "Бухгалтерський кабінет у wireframe орієнтовано на таблиці нарахувань і платежів, фільтри за статусом і зручні дії для підтвердження або перегляду операції. Адміністративний кабінет, своєю чергою, містить окремі блоки для студентів, персоналу і довідників. Усі ці екрани мають спільну логіку: дані подаються через картки і таблиці, а дії виконуються через форми, модальні вікна та кнопки швидких операцій."
    )
    style_body(sec23_p4)
    sec23_p5 = insert_after(
        sec23_p4,
        "Wireframe-підхід корисний ще й тому, що дозволяє перевірити розташування елементів до розроблення стилів. Якщо користувач швидко розуміє, де розміщено навігацію, де відкривається форма, а де показується результат дії, то сам інтерфейс сприймається як логічний і передбачуваний. Саме тому макети стали окремим етапом проектування, а не просто ескізом для оформлення."
    )
    style_body(sec23_p5)
    add_picture_before(find_paragraph(doc, "2.4 Структура сайту/застосунку", 1), wireframe_path)
    wireframe_caption = insert_after(
        find_paragraph(doc, "2.4 Структура сайту/застосунку", 1)._p.getprevious().getparent().paragraphs[-1]
        if False
        else find_paragraph(doc, "2.4 Структура сайту/застосунку", 1),
        "Рисунок 2.8 - Wireframe користувацького інтерфейсу",
    )
    style_caption(wireframe_caption)

    # Section 2.2 - database design.
    sec22 = find_paragraph(doc, "2.2 Проєктування бази даних", 1)
    sec22_p1 = insert_after(
        sec22,
        "Для системи гуртожитку реляційна база даних є природним вибором, оскільки тут фіксуються чітко структуровані дані: користувачі, кімнати, заявки, перепустки, нарахування, платежі, сесії та журнали подій. PostgreSQL добре підходить для такого сценарію завдяки підтримці зовнішніх ключів, транзакцій, індексів і надійного контролю цілісності. Це особливо важливо для моделі, у якій помилки в записах можуть впливати на доступ до приміщень або фінансові операції."
    )
    style_body(sec22_p1)
    sec22_p2 = insert_after(
        sec22_p1,
        "Центральним елементом логічної моделі є зв'язок між ролями та користувачами. Таблиця Role задає тип доступу, таблиця User зберігає персональні дані, а Room описує фізичну кімнату, до якої може бути прив'язаний мешканець. У блоці експлуатації кімнати працюють сутності RelocationRequest і Violation, які зберігають історію переміщень та дисциплінарних подій. У фінансовому блоці розташовані StudentCharge, TariffPlan і Payment, а в сервісному - RepairTicket, GuestPass, PassLog і FileAsset."
    )
    style_body(sec22_p2)
    sec22_p3 = insert_after(
        sec22_p2,
        "Таке групування сутностей відображає реальну організацію життєвого циклу даних у гуртожитку. Студент отримує кімнату, створює заявки, формує перепустки та здійснює платежі; персонал переглядає і змінює стани пов'язаних процесів; адміністратор контролює довідники і доступи. База даних повинна зберігати не лише актуальний стан, а й історію змін, тому окремі таблиці аудиту і сесій доповнюють основну предметну модель."
    )
    style_body(sec22_p3)
    sec22_p4 = insert_after(
        sec22_p3,
        "Нормалізація до третьої нормальної форми використана для усунення дублювання і залежностей, які не визначаються ключем. Наприклад, роль не зберігається як текстове повторення в кожному записі користувача, а винесена в окрему таблицю; кімната описується один раз, а в таблицях заявок і користувачів лише посилаються її ідентифікатор; довідники типу статусів і пріоритетів передаються як обмежені значення або зафіксовані переліки. Завдяки цьому схема стає стійкою до змін і менш вразливою до аномалій оновлення."
    )
    style_body(sec22_p4)
    sec22_p5 = insert_after(
        sec22_p4,
        "Під час фізичного проєктування враховані обмеження полів і правила видалення. Унікальними зроблено email користувача, номер кімнати, код перепустки і хеш токена, а зовнішні ключі мають різну поведінку при видаленні: деякі записи забороняють каскадне стирання, інші зберігають історію або переводять посилання в null. Такий підхід важливий для дипломного проєкту, оскільки тут критичні дані не повинні зникати разом із операційною дією користувача."
    )
    style_body(sec22_p5)
    sec22_p6 = insert_after(
        sec22_p5,
        "ER-діаграма в цьому розділі виконує не декоративну, а методичну функцію: вона наочно демонструє, як сутності пов'язані між собою і які дані можна отримати без складних надлишкових запитів. Саме на основі цієї моделі у подальшому формуються таблиці, мапінг Entity Framework Core, міграції та запити для відображення статистики в інтерфейсі. Тому база даних у проєкті є не другорядним технічним шаром, а основою всієї інформаційної системи."
    )
    style_body(sec22_p6)
    add_picture_before(find_paragraph(doc, "2.3 Макети користувацького інтерфейсу", 1), class_path)
    class_caption = insert_after(
        find_paragraph(doc, "2.3 Макети користувацького інтерфейсу", 1)._p.getprevious().getparent().paragraphs[-1]
        if False
        else find_paragraph(doc, "2.3 Макети користувацького інтерфейсу", 1),
        "Рисунок 2.5 - Діаграма класів предметної області",
    )
    style_caption(class_caption)
    sequence_paragraph = add_picture_before(
        find_paragraph(doc, "2.3 Макети користувацького інтерфейсу", 1), sequence_path
    )
    sequence_caption = insert_after(
        sequence_paragraph,
        "Рисунок 2.6 - Діаграма послідовності взаємодії компонентів",
    )
    style_caption(sequence_caption)

    # Section 2.1 - system functionality and diagrams.
    sec21 = find_paragraph(doc, "2.1 Проєктування функціоналу та структури системи керування гуртожитком", 1)
    sec21_p1 = insert_after(
        sec21,
        "Проєктування системи керування гуртожитком починається з розуміння того, які саме процеси вона має охоплювати. У межах цього проєкту система об'єднує в одному вебзастосунку роботу студентів, коменданта, майстра, охорони, бухгалтера та адміністратора. Кожна роль виконує власний набір дій, проте всі вони працюють з одними й тими ж даними, тому на першому етапі необхідно чітко визначити межі системи та напрямки обміну інформацією."
    )
    style_body(sec21_p1)
    sec21_p2 = insert_after(
        sec21_p1,
        "Основні сценарії роботи в системі охоплюють авторизацію, перегляд стану кімнати, подачу ремонтних і переселенських заявок, оформлення гостьових перепусток, перевірку доступу на вході, нарахування платежів і роботу з довідниками. Саме ці сценарії визначають функціональну декомпозицію, а також показують, як система переходить від однієї ролі до іншої без втрати цілісності даних. Для дипломної роботи важливо не просто перелічити можливості, а показати їхню логічну послідовність."
    )
    style_body(sec21_p2)
    sec21_p3 = insert_after(
        sec21_p2,
        "Контекстна IDEF0-модель потрібна для фіксації меж системи. У ролі входів тут виступають дані про користувачів, кімнати, заявки, платежі та перепустки, у ролі керуючих впливів - правила проживання, ролі доступу, часові обмеження й політики безпеки, у ролі механізмів - вебзастосунок, серверний API, база даних і файлове сховище, а виходом є оновлена інформація, повідомлення та звіти. Такий опис дозволяє побачити, що система не існує окремо від предметної області, а працює як інструмент керування реальними процесами гуртожитку."
    )
    style_body(sec21_p3)
    sec21_p4 = insert_after(
        sec21_p3,
        "DFD-модель, на відміну від IDEF0, підкреслює переміщення даних між учасниками процесу. У нашому випадку це шлях від форми на клієнті до контролера, далі до сервісного шару, після чого дані зберігаються в базі і повертаються в інтерфейс у вигляді таблиць, статусів або попереджень. Наприклад, заявка на переселення проходить перевірку кімнат, валідацію заповнення, запис до БД і відображення у кабінетах студента та коменданта; подібним чином рухається і перевірка гостьової перепустки."
    )
    style_body(sec21_p4)
    sec21_p5 = insert_after(
        sec21_p4,
        "UML-діаграми дають змогу деталізувати поведінку системи на рівні ролей і об'єктів. Діаграма варіантів використання показує межі доступу для кожного користувача; діаграма діяльності відображає сценарії, де є рішення, перевірки й альтернативні гілки; діаграма класів формує уявлення про доменні сутності; діаграма послідовності демонструє взаємодію фронтенду, API, сервісу і бази даних. Усе це разом створює повний ланцюг від бізнес-потреби до технічної реалізації."
    )
    style_body(sec21_p5)
    sec21_p6 = insert_after(
        sec21_p5,
        "Для цієї системи особливо корисні сценарії, пов'язані з перевіркою доступу і зміною станів. Під час авторизації користувач отримує сесію і рольовий набір маршрутів; під час створення ремонтної заявки система фіксує кімнату, пріоритет і автора; під час підтвердження платежу змінюється статус фінансової операції; під час роботи з перепусткою перевіряються часові межі і QR-код. Такі правила краще описувати схемами, ніж суцільним текстом, бо вони наочно показують, де саме виникає управлінське рішення."
    )
    style_body(sec21_p6)
    sec21_p7 = insert_after(
        sec21_p6,
        "Комплект діаграм у цьому розділі не дублює один одного, а розкриває систему з різних боків. IDEF0 відповідає на питання, що система робить і що на неї впливає, DFD - як переміщуються дані, UML use case - хто і що може робити, activity - у якій послідовності виконується дія, class - з яких сутностей складається предметна область, sequence - як компоненти обмінюються повідомленнями. Така багаторівнева модель значно зменшує ризик помилок на стадії реалізації."
    )
    style_body(sec21_p7)

    # Captions and tables in chapter 2.
    for prefix, new_text in [
        ("Рисунок 2.1 -", "Рисунок 2.1 - Контекстна IDEF0-модель системи"),
        ("Рисунок 2.2 -", "Рисунок 2.2 - DFD-модель обміну даними"),
        ("Рисунок 2.3 -", "Рисунок 2.3 - Діаграма варіантів використання системи"),
        ("Рисунок 2.4 -", "Рисунок 2.4 - Діаграма діяльності основних сценаріїв"),
        ("Рисунок 2.5 -", "Рисунок 2.5 - Діаграма класів предметної області"),
        ("Рисунок 2.6 -", "Рисунок 2.6 - Діаграма послідовності взаємодії компонентів"),
        ("Рисунок 2.7 -", "Рисунок 2.7 - ER-модель бази даних"),
        ("Рисунок 2.8 -", "Рисунок 2.8 - Wireframe користувацького інтерфейсу"),
    ]:
        caption = find_paragraph(doc, prefix, 1)
        caption.text = new_text
        style_caption(caption)

    for prefix, new_text in [
        ("Таблиця 2.1 -", "Таблиця 2.1 - Перелік моделей та їх призначення"),
        ("Таблиця 2.3 -", "Таблиця 2.3 - Основні сутності логічної моделі"),
    ]:
        table_caption = find_paragraph(doc, prefix, 1)
        table_caption.text = new_text
        style_caption(table_caption)

    doc.save(DOCX_PATH)
    shutil.copy2(DOCX_PATH, OUTPUT_COPY)
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
