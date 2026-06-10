from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
FIGURE_DIR = OUT_DIR / "figures"
OUT_FILE = OUT_DIR / "Diploma_Report_Student_Dormitory.docx"
FONT_REGULAR = Path("C:/Windows/Fonts/DejaVuSans.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/DejaVuSans-Bold.ttf")


def set_margins(section) -> None:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_page_number_header(section) -> None:
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    run = paragraph.add_run("1")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def set_default_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        set_margins(section)


def set_paragraph_format(paragraph, *, first_line_cm: float = 1.25, align=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(first_line_cm)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    if align is not None:
        paragraph.alignment = align


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return None


def add_body_bold(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True


def add_heading_1(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(12)
    fmt.line_spacing = 1.0
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_heading_2(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_centered(doc: Document, text: str, size: int = 14, bold: bool = False, spacing_after: int = 0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(spacing_after)
    fmt.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_spacer(doc: Document, points: int = 12) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(points)
    paragraph.paragraph_format.space_after = Pt(points)


def add_table_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def style_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(font_path), size)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if text_size(draw, candidate, font)[0] <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int] = (35, 35, 35),
    line_spacing: int = 8,
) -> None:
    x0, y0, x1, y1 = box
    max_width = x1 - x0 - 24
    lines = wrap_text(draw, text, font, max_width)
    line_heights = [text_size(draw, line, font)[1] for line in lines]
    total_height = sum(line_heights) + line_spacing * max(0, len(lines) - 1)
    y = y0 + (y1 - y0 - total_height) // 2
    for index, line in enumerate(lines):
        line_width, line_height = text_size(draw, line, font)
        x = x0 + (x1 - x0 - line_width) // 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height + line_spacing


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    *,
    fill: tuple[int, int, int] = (248, 250, 252),
    outline: tuple[int, int, int] = (59, 130, 246),
    width: int = 4,
    radius: int = 22,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    draw_centered_text(draw, box, text, font)


def draw_ellipse_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    *,
    fill: tuple[int, int, int] = (236, 253, 245),
    outline: tuple[int, int, int] = (16, 185, 129),
    width: int = 4,
) -> None:
    draw.ellipse(box, fill=fill, outline=outline, width=width)
    draw_centered_text(draw, box, text, font)


def draw_diamond(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    *,
    fill: tuple[int, int, int] = (255, 251, 235),
    outline: tuple[int, int, int] = (245, 158, 11),
    width: int = 4,
) -> None:
    x0, y0, x1, y1 = box
    points = [(x0 + (x1 - x0) // 2, y0), (x1, y0 + (y1 - y0) // 2), (x0 + (x1 - x0) // 2, y1), (x0, y0 + (y1 - y0) // 2)]
    draw.polygon(points, fill=fill, outline=outline)
    draw.line(points + [points[0]], fill=outline, width=width)
    draw_centered_text(draw, box, text, font)


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    color: tuple[int, int, int] = (75, 85, 99),
    width: int = 5,
    arrow_size: int = 18,
) -> None:
    draw.line([start, end], fill=color, width=width)
    import math

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    left = (
        end[0] - arrow_size * math.cos(angle - math.pi / 6),
        end[1] - arrow_size * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - arrow_size * math.cos(angle + math.pi / 6),
        end[1] - arrow_size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def create_canvas(width: int = 2200, height: int = 1500) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    return image, draw


def add_canvas_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str | None = None) -> None:
    title_font = load_font(30, bold=True)
    subtitle_font = load_font(20)
    width = draw._image.size[0]
    draw.text(((width - text_size(draw, title, title_font)[0]) // 2, 24), title, font=title_font, fill=(17, 24, 39))
    if subtitle:
        subtitle_width = text_size(draw, subtitle, subtitle_font)[0]
        draw.text(((width - subtitle_width) // 2, 72), subtitle, font=subtitle_font, fill=(75, 85, 99))


def add_corner_label(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], *, fill=(59, 130, 246)) -> None:
    font = load_font(18, bold=True)
    padding_x = 14
    padding_y = 8
    lines = wrap_text(draw, text, font, box[2] - box[0] - 28)
    text_height = sum(text_size(draw, line, font)[1] for line in lines) + 6 * max(0, len(lines) - 1)
    draw.rounded_rectangle(box, radius=18, fill=(239, 246, 255), outline=fill, width=3)
    draw_centered_text(draw, box, text, font, fill=(30, 64, 175), line_spacing=6)


def save_image(image: Image.Image, filename: str) -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / filename
    image.save(path, format="PNG")
    return path


def add_picture_paragraph(doc: Document, path: Path, width_cm: float = 15.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_figure_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def create_idef0_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "IDEF0-модель контексту системи", "Управління студентським гуртожитком")
    font = load_font(26, bold=True)
    small_font = load_font(22)

    central = (720, 540, 1480, 920)
    draw_box(draw, central, "A0\nУправління\nстудентським\nгуртожитком", font, fill=(239, 246, 255))

    draw_box(draw, (120, 590, 520, 860), "Заявки\nна поселення,\nпереселення,\nремонт", small_font, fill=(250, 245, 255), outline=(168, 85, 247))
    draw_arrow(draw, (520, 725), (720, 725), color=(107, 114, 128))

    draw_box(draw, (1220, 590, 1980, 860), "Статуси,\nзвіти,\nсповіщення,\nпідтвердження", small_font, fill=(240, 253, 244), outline=(34, 197, 94))
    draw_arrow(draw, (1480, 735), (1220, 735), color=(107, 114, 128))

    draw_box(draw, (730, 160, 1470, 360), "Правила проживання,\nрегламенти,\nрольові політики", small_font, fill=(255, 251, 235), outline=(245, 158, 11))
    draw_arrow(draw, (1100, 360), (1100, 540), color=(107, 114, 128))

    draw_box(draw, (220, 1040, 620, 1280), "Вебдодаток,\nслужбовий персонал,\nбаза даних", small_font, fill=(248, 250, 252), outline=(14, 165, 233))
    draw_arrow(draw, (620, 1160), (720, 1080), color=(107, 114, 128))

    draw_box(draw, (1580, 1040, 1980, 1280), "Персонал гуртожитку,\nохорона,\nбухгалтерія", small_font, fill=(248, 250, 252), outline=(14, 165, 233))
    draw_arrow(draw, (1480, 1080), (1580, 1160), color=(107, 114, 128))

    return save_image(image, "figure_2_1_idef0.png")


def create_dfd_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "DFD-модель обміну даними", "Контекстний рівень інформаційної системи гуртожитку")
    font = load_font(24, bold=True)
    small_font = load_font(22)

    process = (790, 520, 1410, 920)
    draw_box(draw, process, "Процес\nобробки запитів\nта обліку проживання", font, fill=(239, 246, 255))

    student = (110, 220, 530, 420)
    staff = (110, 1020, 530, 1220)
    security = (1630, 220, 2090, 420)
    accountant = (1630, 1020, 2090, 1220)
    database = (840, 1260, 1360, 1440)

    draw_box(draw, student, "Студент", small_font, fill=(250, 245, 255), outline=(168, 85, 247))
    draw_box(draw, staff, "Комендант / майстер", small_font, fill=(250, 245, 255), outline=(168, 85, 247))
    draw_box(draw, security, "Охорона", small_font, fill=(250, 245, 255), outline=(168, 85, 247))
    draw_box(draw, accountant, "Бухгалтерія", small_font, fill=(250, 245, 255), outline=(168, 85, 247))
    draw_box(draw, database, "База даних", small_font, fill=(240, 253, 244), outline=(34, 197, 94))

    draw_arrow(draw, (530, 320), (790, 620), color=(75, 85, 99))
    draw_arrow(draw, (790, 760), (530, 320), color=(75, 85, 99))
    draw_arrow(draw, (530, 1120), (790, 760), color=(75, 85, 99))
    draw_arrow(draw, (1410, 620), (1630, 320), color=(75, 85, 99))
    draw_arrow(draw, (1410, 760), (1630, 1120), color=(75, 85, 99))
    draw_arrow(draw, (1100, 920), (1100, 1260), color=(75, 85, 99))
    draw_arrow(draw, (1100, 1260), (1100, 920), color=(75, 85, 99))

    add_corner_label(draw, "Заявки, оплати,\nдані профілю", (560, 540, 760, 760), fill=(168, 85, 247))
    add_corner_label(draw, "Статуси,\nповідомлення", (1260, 560, 1560, 780), fill=(34, 197, 94))
    add_corner_label(draw, "Перевірка доступу,\nжурнал візитів", (1450, 520, 1950, 760), fill=(168, 85, 247))
    add_corner_label(draw, "Нарахування,\nпідтвердження оплат", (1450, 860, 1960, 1080), fill=(168, 85, 247))

    return save_image(image, "figure_2_1_dfd.png")


def create_use_case_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "UML use case diagram", "Ролі та ключові сценарії взаємодії з системою")
    actor_font = load_font(22, bold=True)
    use_font = load_font(20)

    boundary = (560, 230, 1640, 1280)
    draw.rounded_rectangle(boundary, radius=28, fill=(248, 250, 252), outline=(100, 116, 139), width=4)
    draw.text((670, 260), "Інформаційна система «Студентський гуртожиток»", font=load_font(24, bold=True), fill=(30, 41, 59))

    use_cases = {
        "Авторизація": (820, 360, 1280, 470),
        "Перегляд\nособистої\nінформації": (650, 560, 1000, 700),
        "Створення\nзаявки на ремонт": (1080, 560, 1450, 700),
        "Оформлення\nперепустки": (650, 770, 1000, 910),
        "Облік\nоплат": (1080, 770, 1450, 910),
        "Керування\nзаселенням": (840, 980, 1280, 1120),
    }
    for text, box in use_cases.items():
        draw_ellipse_box(draw, box, text, use_font)

    actors = {
        "Студент": (120, 430, 330, 560),
        "Комендант": (120, 720, 330, 850),
        "Майстер": (1720, 430, 1930, 560),
        "Бухгалтер": (1720, 720, 1930, 850),
    }
    for name, box in actors.items():
        draw_box(draw, box, name, actor_font, fill=(250, 245, 255), outline=(168, 85, 247))

    connections = [
        ((330, 495), (820, 415)),
        ((330, 495), (650, 630)),
        ((330, 785), (650, 840)),
        ((330, 785), (840, 1050)),
        ((1720, 495), (1450, 630)),
        ((1720, 785), (1450, 840)),
        ((1450, 630), (1450, 630)),
    ]
    for start, end in connections:
        if start != end:
            draw_arrow(draw, start, end, color=(75, 85, 99), width=4, arrow_size=14)

    draw_arrow(draw, (330, 495), (820, 415), color=(75, 85, 99), width=4, arrow_size=14)
    draw_arrow(draw, (330, 495), (650, 630), color=(75, 85, 99), width=4, arrow_size=14)
    draw_arrow(draw, (330, 785), (650, 840), color=(75, 85, 99), width=4, arrow_size=14)
    draw_arrow(draw, (330, 785), (840, 1050), color=(75, 85, 99), width=4, arrow_size=14)
    draw_arrow(draw, (1720, 495), (1450, 630), color=(75, 85, 99), width=4, arrow_size=14)
    draw_arrow(draw, (1720, 785), (1450, 840), color=(75, 85, 99), width=4, arrow_size=14)
    draw_arrow(draw, (1720, 785), (1280, 1050), color=(75, 85, 99), width=4, arrow_size=14)

    return save_image(image, "figure_2_1_use_case.png")


def create_activity_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "UML activity diagram", "Створення заявки на ремонт у гуртожитку")
    font = load_font(22, bold=True)
    small_font = load_font(20)

    start = (960, 180, 1240, 270)
    draw_ellipse_box(draw, start, "Початок", small_font, fill=(224, 242, 254), outline=(14, 165, 233))

    steps = {
        "Обрати\nрозділ\n«Ремонтні заявки»": (820, 330, 1380, 500),
        "Заповнити\nформу та додати\nопис проблеми": (780, 560, 1420, 740),
        "Перевірити\nкоректність\nданих": (860, 790, 1340, 940),
        "Зберегти\nзаявку": (900, 1010, 1300, 1160),
    }
    for text, box in steps.items():
        draw_box(draw, box, text, font, fill=(248, 250, 252))

    decision = (880, 1220, 1320, 1400)
    draw_diamond(draw, decision, "Дані\nкоректні?", small_font)

    end_yes = (1450, 1220, 1730, 1310)
    end_no = (250, 1220, 540, 1310)
    draw_ellipse_box(draw, end_yes, "Заявка\nстворена", small_font, fill=(236, 253, 245), outline=(16, 185, 129))
    draw_ellipse_box(draw, end_no, "Повернутись\nдо форми", small_font, fill=(255, 241, 242), outline=(244, 63, 94))

    draw_arrow(draw, (1100, 270), (1100, 330))
    draw_arrow(draw, (1100, 500), (1100, 560))
    draw_arrow(draw, (1100, 740), (1100, 790))
    draw_arrow(draw, (1100, 940), (1100, 1010))
    draw_arrow(draw, (1100, 1160), (1100, 1220))
    draw_arrow(draw, (1320, 1310), (1450, 1265), color=(22, 163, 74))
    draw_arrow(draw, (880, 1310), (540, 1265), color=(220, 38, 38))
    draw.text((1420, 1180), "так", font=small_font, fill=(22, 163, 74))
    draw.text((580, 1180), "ні", font=small_font, fill=(220, 38, 38))
    draw_arrow(draw, (405, 1265), (780, 1265), color=(220, 38, 38))

    return save_image(image, "figure_2_1_activity.png")


def create_class_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "UML class diagram", "Основні доменні сутності та зв'язки")
    class_font = load_font(18, bold=True)
    attr_font = load_font(14)

    boxes = {
        "User": ((120, 260, 520, 460), ["+Id: Guid", "+FullName: string", "+RoleId: Guid", "+Email: string"], (59, 130, 246)),
        "Room": ((720, 220, 1120, 400), ["+Id: Guid", "+RoomNumber: string", "+Capacity: int", "+Floor: int"], (16, 185, 129)),
        "RepairTicket": ((1320, 220, 1760, 440), ["+Id: Guid", "+Title: string", "+Status: TicketStatus", "+RoomId: Guid"], (245, 158, 11)),
        "GuestPass": ((260, 760, 700, 980), ["+Id: Guid", "+GuestFullName: string", "+ValidFrom: DateTime", "+AccessCode: string"], (168, 85, 247)),
        "Payment": ((940, 760, 1380, 960), ["+Id: Guid", "+Amount: decimal", "+Status: string", "+UserId: Guid"], (14, 165, 233)),
        "Role": ((1600, 760, 1980, 920), ["+Id: Guid", "+Name: string"], (34, 197, 94)),
    }

    for title, (box, attrs, outline) in boxes.items():
        draw.rounded_rectangle(box, radius=18, fill=(248, 250, 252), outline=outline, width=4)
        x0, y0, x1, y1 = box
        draw.rectangle((x0, y0, x1, y0 + 42), fill=(226, 232, 240))
        title_width, _ = text_size(draw, title, class_font)
        draw.text((x0 + (x1 - x0 - title_width) // 2, y0 + 10), title, font=class_font, fill=(30, 41, 59))
        y = y0 + 56
        for attr in attrs:
            draw.text((x0 + 18, y), attr, font=attr_font, fill=(51, 65, 85))
            y += 26

    draw_arrow(draw, (520, 350), (720, 310), color=(75, 85, 99))
    draw_arrow(draw, (1120, 300), (1320, 310), color=(75, 85, 99))
    draw_arrow(draw, (940, 420), (520, 760), color=(75, 85, 99))
    draw_arrow(draw, (1120, 400), (1180, 760), color=(75, 85, 99))
    draw_arrow(draw, (1760, 440), (1740, 760), color=(75, 85, 99))

    add_corner_label(draw, "1:N", (590, 280, 690, 350), fill=(59, 130, 246))
    add_corner_label(draw, "1:N", (1180, 260, 1290, 340), fill=(59, 130, 246))
    add_corner_label(draw, "1:N", (860, 520, 990, 600), fill=(59, 130, 246))
    add_corner_label(draw, "1:N", (1320, 540, 1440, 620), fill=(59, 130, 246))
    add_corner_label(draw, "1:N", (1640, 540, 1760, 620), fill=(59, 130, 246))

    return save_image(image, "figure_2_2_class.png")


def create_sequence_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "UML sequence diagram", "Створення заявки на ремонт")

    actor_font = load_font(16, bold=True)
    message_font = load_font(16)
    participants = [
        ("Студент", 180, (168, 85, 247)),
        ("Web UI", 670, (59, 130, 246)),
        ("API", 1120, (16, 185, 129)),
        ("DB", 1580, (245, 158, 11)),
    ]
    top = 180
    bottom = 1300
    for name, x, color in participants:
        draw.rounded_rectangle((x - 110, top, x + 110, top + 70), radius=16, fill=(248, 250, 252), outline=color, width=4)
        width, _ = text_size(draw, name, actor_font)
        draw.text((x - width // 2, top + 24), name, font=actor_font, fill=(30, 41, 59))
        draw.line((x, top + 70, x, bottom), fill=(148, 163, 184), width=3)

    messages = [
        ((290, 330), (560, 330), "1. submitRepairForm()"),
        ((780, 440), (1000, 440), "2. POST /repair-tickets"),
        ((1220, 550), (1500, 550), "3. INSERT ticket"),
        ((1500, 650), (1220, 650), "4. created ticket id"),
        ((1000, 760), (780, 760), "5. success response"),
        ((560, 870), (290, 870), "6. confirmation"),
    ]
    for start, end, label in messages:
        draw_arrow(draw, start, end, color=(75, 85, 99), width=4, arrow_size=12)
        label_width, _ = text_size(draw, label, message_font)
        draw.text(((start[0] + end[0] - label_width) // 2, start[1] - 26), label, font=message_font, fill=(51, 65, 85))

    return save_image(image, "figure_2_2_sequence.png")


def create_wireframe_figure() -> Path:
    image = Image.new("RGB", (2200, 1400), (250, 250, 249))
    draw = ImageDraw.Draw(image)
    add_canvas_title(draw, "Wireframe інтерфейсу", "Головна сторінка рольового кабінету")
    title_font = load_font(20, bold=True)
    text_font = load_font(16)

    draw.rounded_rectangle((120, 140, 2080, 1260), radius=28, outline=(100, 116, 139), width=4, fill=(255, 255, 255))
    draw.rectangle((120, 140, 2080, 220), fill=(241, 245, 249))
    draw.text((170, 165), "Студентський гуртожиток", font=title_font, fill=(15, 23, 42))
    draw.rounded_rectangle((160, 280, 520, 1180), radius=24, fill=(248, 250, 252), outline=(148, 163, 184), width=3)
    draw.text((220, 320), "Навігація", font=title_font, fill=(30, 41, 59))
    for idx, item in enumerate(["Дашборд", "Заявки", "Перепустки", "Фінанси", "Користувачі", "Звіти"]):
        y = 390 + idx * 100
        draw.rounded_rectangle((200, y, 480, y + 60), radius=16, fill=(239, 246, 255), outline=(191, 219, 254), width=2)
        draw.text((230, y + 18), item, font=text_font, fill=(37, 99, 235))

    draw.rounded_rectangle((580, 280, 2020, 420), radius=24, fill=(248, 250, 252), outline=(148, 163, 184), width=3)
    draw.text((630, 320), "Пошук, фільтри, швидкі дії", font=title_font, fill=(30, 41, 59))

    card_positions = [(580, 480), (1030, 480), (1480, 480), (580, 780), (1030, 780), (1480, 780)]
    card_titles = ["Активні заявки", "Заселеність", "Борг", "Нові повідомлення", "Пропуски", "Платежі"]
    for (x, y), title in zip(card_positions, card_titles):
        draw.rounded_rectangle((x, y, x + 380, y + 220), radius=24, fill=(255, 255, 255), outline=(226, 232, 240), width=3)
        draw.text((x + 24, y + 24), title, font=title_font, fill=(30, 41, 59))
        draw.rectangle((x + 24, y + 84, x + 260, y + 116), fill=(219, 234, 254))
        draw.rectangle((x + 24, y + 140, x + 320, y + 170), fill=(241, 245, 249))

    return save_image(image, "figure_2_4_wireframe.png")
def create_er_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "ER-модель бази даних", "Основні сутності та зв'язки домену гуртожитку")
    entity_font = load_font(20, bold=True)
    small_font = load_font(18)

    entities = {
        "Roles": (120, 260, 420, 390, "1"),
        "Users": (720, 220, 1100, 390, "N"),
        "Rooms": (1370, 220, 1750, 390, "N"),
        "RepairTickets": (360, 650, 780, 830, "N"),
        "GuestPasses": (960, 620, 1380, 800, "N"),
        "Payments": (1560, 640, 1980, 820, "N"),
        "AuditLogs": (700, 1080, 1120, 1260, "N"),
    }

    for name, (x0, y0, x1, y1, mult) in entities.items():
        draw_box(draw, (x0, y0, x1, y1), name, entity_font, fill=(248, 250, 252))
        draw.text((x1 - 52, y1 - 42), mult, font=load_font(24, bold=True), fill=(17, 24, 39))

    draw_arrow(draw, (420, 325), (720, 305), color=(75, 85, 99))
    draw_arrow(draw, (1100, 305), (1370, 305), color=(75, 85, 99))
    draw_arrow(draw, (930, 390), (560, 650), color=(75, 85, 99))
    draw_arrow(draw, (1100, 390), (1160, 620), color=(75, 85, 99))
    draw_arrow(draw, (1420, 390), (1650, 640), color=(75, 85, 99))
    draw_arrow(draw, (910, 830), (910, 1080), color=(75, 85, 99))
    draw_arrow(draw, (1160, 800), (910, 1080), color=(75, 85, 99))

    add_corner_label(draw, "1:M", (520, 250, 660, 330), fill=(59, 130, 246))
    add_corner_label(draw, "1:M", (1180, 250, 1320, 330), fill=(59, 130, 246))
    add_corner_label(draw, "1:M", (530, 560, 680, 640), fill=(59, 130, 246))
    add_corner_label(draw, "1:M", (1180, 560, 1320, 640), fill=(59, 130, 246))
    add_corner_label(draw, "1:M", (1540, 570, 1680, 650), fill=(59, 130, 246))
    add_corner_label(draw, "1:M", (770, 930, 910, 1010), fill=(59, 130, 246))

    return save_image(image, "figure_2_2_er.png")


def add_toc_placeholder(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("ЗМІСТ").bold = True
    paragraph.runs[0].font.name = "Times New Roman"
    paragraph.runs[0].font.size = Pt(14)

    items = [
        ("ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ, СКОРОЧЕНЬ І ТЕРМІНІВ", "4"),
        ("ВСТУП", "5"),
        ("РОЗДІЛ 1 ПОСТАНОВКА ЗАДАЧІ", "7"),
        ("1.1 Аналіз предметної області (опис бізнес-процесу)", "7"),
        ("1.2 Аналіз існуючих рішень", "10"),
        ("1.3 Формування вимог до проєкту", "12"),
        ("РОЗДІЛ 2 ПРОЄКТУВАННЯ СИСТЕМИ", "15"),
        ("2.1 Діаграми IDEF0, Data Flow, UML-моделі", "15"),
        ("2.2 ER-модель бази даних і нормалізація до 3 нормальної форми", "17"),
        ("2.3 Макет користувацького інтерфейсу", "20"),
        ("2.4 Структура сайту", "22"),
        ("2.5 Алгоритми, що використовувались", "23"),
        ("РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ", "25"),
        ("3.1 Обґрунтування вибору стеку розробки", "25"),
        ("3.2 Реалізація бази даних", "27"),
        ("3.3 Дизайн користувацького інтерфейсу", "29"),
        ("3.4 Розроблення заходів для захисту інформації", "31"),
        ("РОЗДІЛ 4 ТЕСТУВАННЯ ТА ВПРОВАДЖЕННЯ", "33"),
        ("4.1 Обґрунтування вибору методів тестування", "33"),
        ("4.2 Формування тест-плану", "34"),
        ("4.3 Впровадження системи", "36"),
        ("ВИСНОВКИ", "38"),
        ("СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", "39"),
        ("ДОДАТКИ", "41"),
    ]
    for title, page in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(f"{title}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        if title.startswith("1.") or title.startswith("2.") or title.startswith("3.") or title.startswith("4."):
            paragraph.paragraph_format.left_indent = Cm(1)
        elif title in {"ВСТУП", "ВИСНОВКИ", "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", "ДОДАТКИ"}:
            run.bold = True
        if page:
            paragraph.add_run(" " * 6 + page)


def add_reference(doc: Document, text: str) -> None:
    add_body(doc, text)


def add_title_page_1(doc: Document) -> None:
    lines = [
        "МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ",
        'НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ «ПОЛТАВСЬКА ПОЛІТЕХНІКА ІМЕНІ ЮРІЯ КОНДРАТЮКА»',
        "Навчально-науковий інститут інформаційних технологій та робототехніки",
        "Кафедра комп’ютерних та інформаційних технологій і систем",
    ]
    for line in lines:
        add_centered(doc, line, size=14, bold=True)

    add_spacer(doc, 30)
    add_centered(doc, "КВАЛІФІКАЦІЙНА РОБОТА БАКАЛАВРА", size=14, bold=True)
    add_centered(doc, 'спеціальність 122 «Комп’ютерні науки»', size=14)
    add_spacer(doc, 12)
    add_centered(doc, "на тему", size=14)
    add_centered(doc, '«Модель інформаційної системи "Студентський гуртожиток"»', size=16, bold=True)

    add_spacer(doc, 36)
    add_centered(doc, "Виконав: студент групи 404-ТН", size=14)
    add_centered(doc, "Чорнорук Роман Олександрович", size=14, bold=True)
    add_spacer(doc, 12)
    add_centered(doc, "Керівник роботи: ____________________________", size=14)
    add_centered(doc, "Рецензент: _________________________________", size=14)
    add_centered(doc, "Завідувач кафедри: _________________________", size=14)
    add_spacer(doc, 18)
    add_centered(doc, "Полтава - 2026", size=14)


def add_title_page_2(doc: Document) -> None:
    add_centered(doc, "ЗАВДАННЯ", size=16, bold=True)
    add_centered(doc, "НА КВАЛІФІКАЦІЙНУ РОБОТУ БАКАЛАВРА", size=14, bold=True)
    add_spacer(doc, 12)
    add_body_bold(doc, 'Тема роботи: «Модель інформаційної системи "Студентський гуртожиток"»')
    add_body(doc, "Мета роботи полягає у створенні моделі веборієнтованої інформаційної системи для підтримки ключових процесів студентського гуртожитку.")
    add_body(doc, "У межах роботи передбачено аналіз предметної області, проєктування логічної та фізичної моделі даних, розроблення користувацького інтерфейсу, опис алгоритмів, реалізацію програмної частини та підготовку матеріалів для тестування і впровадження.")
    add_body(doc, "Основні результати повинні включати: модель бізнес-процесів, ER-модель бази даних, прототип інтерфейсу, опис стеку розробки, схему розгортання в контейнерах та комплект додатків із тестовими матеріалами.")
    add_spacer(doc, 12)
    add_body(doc, "Студент _______________________    Керівник _______________________")
    add_body(doc, "Дата видачі завдання _______________________")
    add_spacer(doc, 18)
    add_body(doc, "Календарний план виконання роботи подано у складі завдання та може бути уточнений керівником під час підготовки остаточної версії документа.")


def build_document() -> Document:
    doc = Document()
    set_default_styles(doc)

    # Title page 1.
    section = doc.sections[0]
    set_margins(section)
    add_title_page_1(doc)

    # Title page 2.
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(section)
    add_title_page_2(doc)

    # Main body with page numbers starting from 2 as requested.
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(section)
    set_page_number_start(section, 2)
    add_page_number_header(section)

    # Summary / abstract.
    add_heading_1(doc, "РЕФЕРАТ")
    add_body(doc, "Кваліфікаційна робота бакалавра присвячена моделюванню інформаційної системи для студентського гуртожитку, що поєднує в єдиному цифровому середовищі облік мешканців, кімнатного фонду, заявок на ремонт, гостьових перепусток, нарахувань за проживання та проходження контролю доступу.")
    add_body(doc, "Об’єктом розробки є сукупність адміністративних, побутових і сервісних процесів гуртожитку, а предметом - методи та засоби проєктування веборієнтованої інформаційної системи, логічної моделі даних і інтерфейсу користувача.")
    add_body(doc, "Метою роботи є створення узгодженої моделі інформаційної системи, яка забезпечує рольовий доступ, швидку обробку звернень, контроль фінансових операцій і документування дій користувачів. Для досягнення мети використано методи аналізу предметної області, побудови IDEF0- та DFD-моделей, UML-моделювання, ER-проєктування, нормалізації даних, прототипування інтерфейсу та аналізу реалізації програмного забезпечення.")
    add_body(doc, "Отримані результати включають опис бізнес-процесів гуртожитку, набір функціональних і нефункціональних вимог, проєкт реляційної бази даних, структуру вебсайту, моделі основних сценаріїв взаємодії, обґрунтування вибору стеку розробки, опис заходів захисту інформації та схему контейнерного розгортання.")
    add_body(doc, "Ключові слова: інформаційна система, студентський гуртожиток, вебдодаток, база даних, UML, DFD, IDEF0, React, ASP.NET Core, PostgreSQL, Docker, RBAC, JWT.")

    doc.add_page_break()
    add_toc_placeholder(doc)

    doc.add_page_break()
    add_heading_1(doc, "ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ, СКОРОЧЕНЬ І ТЕРМІНІВ")
    abbreviations = [
        ("ІС", "інформаційна система"),
        ("БД", "база даних"),
        ("API", "інтерфейс прикладного програмування"),
        ("SPA", "односторінковий вебдодаток"),
        ("UML", "Unified Modeling Language"),
        ("DFD", "Data Flow Diagram"),
        ("IDEF0", "функціональна модель бізнес-процесу"),
        ("ER", "Entity-Relationship модель"),
        ("RBAC", "рольовий контроль доступу"),
        ("JWT", "JSON Web Token"),
        ("CSRF", "Cross-Site Request Forgery"),
        ("CRUD", "створення, читання, оновлення, видалення"),
        ("DBMS", "система керування базою даних"),
    ]
    for term, meaning in abbreviations:
        add_body(doc, f"{term} - {meaning}.")

    doc.add_page_break()
    add_heading_1(doc, "ВСТУП")
    intro_paragraphs = [
        "Актуальність теми зумовлена тим, що студентський гуртожиток є середовищем із великою кількістю паралельних процесів: заселення і переселення мешканців, контроль технічного стану кімнат, облік гостьових візитів, нарахування платежів, фіксація порушень та оперативна взаємодія між студентами і персоналом. За паперового або фрагментованого способу ведення цих процесів зростає ризик втрати інформації, дублювання записів і затримок у прийнятті рішень.",
        "Перехід до моделі інформаційної системи дає змогу не лише автоматизувати окремі операції, а й узгодити їх у межах єдиного інформаційного простору. Це особливо важливо для гуртожитку, де різні ролі мають власні повноваження і бачать лише ту інформацію, яка потрібна їм для виконання службових обов’язків. Такий підхід підвищує прозорість обліку, прискорює реагування на побутові проблеми та спрощує контроль доступу.",
        "Метою роботи є розроблення моделі інформаційної системи \"Студентський гуртожиток\", здатної підтримувати рольову взаємодію користувачів, структурувати основні дані предметної області та забезпечити практичну основу для подальшої програмної реалізації. Для досягнення мети необхідно проаналізувати предметну область, сформувати вимоги, спроєктувати моделі процесів і даних, обґрунтувати вибір технологічного стеку та описати впровадження системи в контейнерному середовищі.",
        "Об’єктом розробки є інформаційні процеси, що супроводжують діяльність студентського гуртожитку, а предметом - методи моделювання, проєктування та реалізації веборієнтованої системи керування цими процесами. Очікуваним результатом є цілісна архітектурна і логічна модель, придатна для впровадження у вигляді реального програмного продукту.",
        "Практичне значення роботи полягає в тому, що запропонована модель може бути використана як основа для створення внутрішньої інформаційної системи гуртожитку, яка об’єднає адміністративні, сервісні та фінансові сценарії в одному інтерфейсі. Окрему цінність становить можливість подальшого розширення системи без перебудови її базової структури.",
    ]
    for paragraph in intro_paragraphs:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading_1(doc, "РОЗДІЛ 1 ПОСТАНОВКА ЗАДАЧІ")
    add_heading_2(doc, "1.1 Аналіз предметної області (опис бізнес-процесу)")
    part_1_1 = [
        "Предметна область студентського гуртожитку охоплює взаємопов’язані процеси, що супроводжують повсякденне проживання мешканців і службову діяльність персоналу. До базових учасників належать студент, комендант, майстер, охоронець, бухгалтер та адміністратор. Кожна з ролей має власний набір дій і доступу до даних, тому модель системи повинна передбачати чітке розмежування прав.",
        "Основний бізнес-процес починається з авторизації користувача і визначення його ролі. Після входу студент отримує доступ до власного кабінету, де може переглянути стан проживання, подати заявку на ремонт, створити запит на переселення, оформити гостьову перепустку та перевірити фінансові нарахування. Комендант опрацьовує заявки на переселення, контролює заселеність кімнат і фіксує дисциплінарні зауваження. Майстер переглядає ремонтні заявки, змінює їх статус і додає службові коментарі. Охоронець перевіряє перепустки гостей, а бухгалтер підтверджує платежі та контролює заборгованість.",
        "У межах предметної області важливими є дані про кімнати, поверхи, місткість, вартість проживання, фактичну зайнятість, стан ремонту, список мешканців і історію подій. Саме ці дані формують основу для розрахунку доступних місць, перевірки можливості переселення і прийняття рішень щодо заселення нових студентів. Окремо зберігаються записи про фінансові нарахування, платежі та статус сплати, оскільки ці відомості потрібні як студенту, так і персоналу.",
        "У реальному житті частина процесів виконується вручну або за допомогою розрізнених таблиць і журналів. Це створює затримки під час погодження переселення, ускладнює пошук інформації про заявки та робить контроль відвідувань менш надійним. Отже, для даної предметної області доцільно створити єдину інформаційну систему, яка централізує дані та забезпечує швидкий доступ до них залежно від ролі користувача.",
        "На рівні бізнес-логіки систему можна подати як послідовність взаємодій між користувачем, серверною частиною, базою даних і зовнішніми службами. Результатом кожної дії є зміна стану конкретної сутності: заявки, кімнати, перепустки, платежу або дисциплінарного запису. Така модель дає змогу пов’язати між собою побутові, адміністративні та фінансові аспекти роботи гуртожитку.",
    ]
    for paragraph in part_1_1:
        add_body(doc, paragraph)

    add_heading_2(doc, "1.2 Аналіз існуючих рішень")
    part_1_2 = [
        "Для формування вимог до власної системи доцільно проаналізувати кілька поширених рішень для керування студентським житлом і суміжними процесами. У межах порівняння доречно взяти StarRez, Entrata Student, RMS Cloud та Roompact, оскільки ці продукти покривають близькі, але не ідентичні сценарії: від класичного обліку гуртожитків до підтримки residence life, аналітики й автоматизації комунікації.",
        "Під час аналізу важливо дивитися не лише на наявність окремих функцій, а й на те, наскільки система відповідає потребам локального студентського гуртожитку. Для такої задачі критичними є функціональне покриття, зручність інтерфейсу, чітка рольова модель, інтеграції з іншими системами закладу, наявність аналітики або елементів штучного інтелекту, вартість впровадження та обмеження, які можуть ускладнити практичне використання.",
        "StarRez є однією з найбільш функціонально насичених платформ у цій ніші. На офіційних матеріалах системи зазначаються модулі housing management, billing, maintenance, visitor tracking, analytics та StarRez Intelligence, який підтримує AI-assisted workflows і пошук тенденцій у даних. Це робить продукт сильним з позиції комплексності, але для невеликого або середнього гуртожитку він може бути надмірно складним і дорогим у впровадженні.",
        "Entrata Student орієнтується на student housing і робить акцент на roommate matching, move-in та turn management, bulk assignment board, renewals і ширшій AI-платформі ELI+. Система має виразну автоматизацію й розвинуті інтеграції з іншими сервісами, однак її логіка ближча до великого комерційного student housing, ніж до локальної моделі гуртожитку з дисциплінарним контролем, побутовими зверненнями та вузькою адміністративною структурою.",
        "RMS Cloud пропонує reservations management, maintenance, revenue management та business intelligence, а також API access і партнерські інтеграції. Це сильне рішення для гнучкого управління об’єктами розміщення, але його ядро побудоване навколо hospitality-логіки, тому воно краще підходить для готельних і змішаних сервісних сценаріїв, ніж для навчального гуртожитку.",
        "Roompact займає іншу нішу: система фокусується на residence life, формах, duty, events, agreements, insights і повідомленнях. Вона добре доповнює класичну housing management platform, має прозору модель оплати та підтримує інтеграції через HMS/SIS, SSO й експортний JSON API. Водночас Roompact не є повноцінною системою обліку гуртожитку, а радше інструментом для освітньої та комунікативної роботи всередині вже існуючої інфраструктури.",
        "Отже, аналіз показує, що для дипломної роботи доцільно проєктувати власну систему, яка поєднає базові функції обліку проживання, ремонтів, перепусток, фінансів і ролей, але буде простішою та ближчою до конкретних процесів студентського гуртожитку. Саме тому в проєкті варто взяти кращі практики існуючих платформ, але не копіювати їхню комерційну надлишковість.",
    ]
    for paragraph in part_1_2:
        add_body(doc, paragraph)

    add_table_caption(doc, "Таблиця 1.3 - Порівняння існуючих рішень")
    add_table(
        doc,
        [
            "Система",
            "Функціональність",
            "Зручність інтерфейсу",
            "Рольова модель",
            "Інтеграції",
            "ШІ / аналітика",
            "Вартість",
            "Обмеження",
        ],
        [
            [
                "StarRez",
                "Комплексний облік гуртожитків, billing, maintenance, visitor tracking, roommate matching.",
                "Професійний інтерфейс, але насичений і вимогливий до навчання персоналу.",
                "Сильне розмежування ролей resident / staff / admin.",
                "Партнерська мережа інтеграцій, у т.ч. з корпоративними сервісами.",
                "StarRez Intelligence, dashboards, reporting та прогнозування тенденцій.",
                "Ціна не опублікована, доступна за запитом.",
                "Надлишкова складність і висока вартість для невеликого гуртожитку.",
            ],
            [
                "Entrata Student",
                "Roommate matching, move-in checklists, bulk assignment, renewals, turn management.",
                "Сучасний workflow-інтерфейс, орієнтований на операторів великого фонду.",
                "Рольова модель з опорою на resident та operator сценарії.",
                "Розвинуті інтеграції, у т.ч. Student.com, StudentRent та інші партнери.",
                "ELI+ та AI-powered workflows, leasing AI, maintenance AI, analytics.",
                "Публічної ціни немає, продаж через демонстрацію / комерційний контакт.",
                "Більше відповідає великому student housing, ніж локальному гуртожитку ЗВО.",
            ],
            [
                "RMS Cloud",
                "Reservations, maintenance, revenue management, business intelligence, rate management.",
                "Функціональний, але hospitality-орієнтований інтерфейс.",
                "Ролі типові для об’єктів розміщення; менш спеціалізовані для гуртожитку.",
                "Partners and integrations, API access, channel integrations.",
                "Аналітичні панелі, forecasting, dynamic pricing; явного AI-фокусу менше.",
                "Ціна за запитом, є тарифні плани і add-ons.",
                "Сильний у гостьових та готельних сценаріях, менш точний для побуту гуртожитку.",
            ],
            [
                "Roompact",
                "Residence life, forms, duty, events, agreements, messaging, insights.",
                "Дуже зручний, візуально зрозумілий, орієнтований на staff/resident взаємодію.",
                "Розмежування staff, residents, admins та доступів до функцій.",
                "HMS/SIS інтеграції, SSO, automated data feeds, outbound JSON API.",
                "Trends, Insights, Wolfie AI та AI-powered data analysis.",
                "Прозора щорічна оплата за кількістю мешканців + setup fee.",
                "Не замінює housing management software і залежить від зовнішньої HMS.",
            ],
        ],
        widths=[2.2, 3.1, 2.5, 2.5, 2.4, 2.2, 1.8, 2.8],
    )

    add_body(
        doc,
        "Порівняння показує, що жодна з розглянутих систем не покриває одночасно побутові, фінансові, дисциплінарні та пропускні сценарії саме в логіці навчального гуртожитку без значної адаптації. Це підтверджує доцільність розроблення власної інформаційної системи, орієнтованої на конкретну предметну область і обмежений набір дій користувачів.",
    )

    add_heading_2(doc, "1.3 Формування вимог до проєкту")
    part_1_3 = [
        "На основі аналізу предметної області та існуючих рішень формуються функціональні й нефункціональні вимоги до системи. Функціональна частина повинна забезпечувати автентифікацію користувачів, розмежування доступу за ролями, ведення кімнатного фонду, роботу із заявками на ремонт і переселення, створення та перевірку гостьових перепусток, облік платежів і нарахувань, а також керування довідниками системи.",
        "Особлива увага приділяється підтримці ролей. Студент має бачити тільки власні дані, комендант - інформацію про заселеність і дисциплінарні події, майстер - активні ремонтні заявки, охоронець - перепустки і журнал проходу, бухгалтер - оплату і борг, а адміністратор - користувачів, персонал та системні довідники. Такий підхід зменшує ризики несанкціонованого доступу та робить інтерфейс зрозумілішим для кожної групи користувачів.",
        "До нефункціональних вимог належать безпека, надійність, продуктивність, підтримуваність і зручність використання. Система має працювати у веббраузері, швидко реагувати на дії користувача, коректно поводитися при одночасній роботі кількох ролей і забезпечувати автоматичне завершення неактивних сесій. Окремо слід передбачити журналювання критичних подій і контроль цілісності даних.",
        "Результати формування вимог подано у таблицях 1.1 та 1.2, де узагальнено функціональні й нефункціональні характеристики майбутньої системи. Саме ці вимоги визначають межі подальшого проєктування і стають основою для розроблення моделей у наступному розділі.",
    ]
    add_table_caption(doc, "Таблиця 1.1 - Функціональні вимоги до системи")
    add_table(
        doc,
        ["Підсистема", "Ключові можливості"],
        [
            ["Користувачі та ролі", "Автентифікація, рольовий доступ, персоналізований інтерфейс, керування обліковими записами."],
            ["Кімнатний фонд", "Ведення довідника кімнат, поверхів, місткості, вартості проживання та стану ремонту."],
            ["Переселення", "Подання заявки, перегляд доступних місць, погодження або відхилення запиту, оновлення кімнати."],
            ["Ремонтні заявки", "Створення звернення, зміна статусу, коментар майстра, прикріплення файлів."],
            ["Гостьові перепустки", "Оформлення перепустки, формування коду доступу, перевірка на вході та виході."],
            ["Фінанси", "Нарахування, облік платежів, підтвердження оплати, відображення заборгованості."],
            ["Адміністрування", "Керування довідниками, тарифами, категоріями поломок, системними ролями."],
        ],
        widths=[5, 11],
    )
    add_table_caption(doc, "Таблиця 1.2 - Нефункціональні вимоги до системи")
    add_table(
        doc,
        ["Категорія", "Вимога"],
        [
            ["Безпека", "Зберігання паролів у хешованому вигляді, контроль сесій, CSRF-захист, рольові політики доступу."],
            ["Продуктивність", "Швидке завантаження основних сторінок, робота без помітних затримок при типовому навантаженні."],
            ["Надійність", "Автоматичне резервування критичних даних і стійкість до втрати окремих сесій чи запитів."],
            ["Зручність", "Адаптивний вебінтерфейс, логічна структура сторінок, мінімальна кількість кроків для основних дій."],
            ["Підтримуваність", "Модульна архітектура, відокремлення шарів, можливість розширення без перебудови системи."],
        ],
        widths=[4, 12],
    )
    for paragraph in part_1_3:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading_1(doc, "РОЗДІЛ 2 ПРОЄКТУВАННЯ СИСТЕМИ")
    add_heading_2(doc, "2.1 Діаграми IDEF0 та Data Flow")
    part_2_1 = [
        "На етапі проєктування використано дві базові моделі опису процесів. IDEF0 дає змогу описати систему на рівні функцій і підкреслити зв’язки між входами, керуючими впливами, механізмами виконання та виходами. У контекстній моделі системи входами виступають дані про користувачів, кімнати, заявки, перепустки й платежі, керуванням - правила проживання та політики доступу, механізмами - вебдодаток, сервер і база даних, а виходами - оновлені записи, звіти і повідомлення.",
        "Діаграми потоків даних використовуються для відображення переміщення інформації між користувачем, інтерфейсом, API, базою даних і зовнішніми сервісами. На рівні контексту система приймає запити від студентів і персоналу, передає їх до серверної логіки та повертає результати у вигляді сторінок, таблиць і статусів виконання. Для таких сценаріїв особливо важливі шляхи створення заявки, погодження переселення, валідації перепустки і підтвердження платежу.",
        "У сукупності ці діаграми формують цілісне уявлення про систему і зменшують ризик помилок на етапі реалізації. Саме тому вони повинні бути розміщені у графічній частині або в додатках до пояснювальної записки як обов’язковий ілюстративний матеріал.",
    ]
    for paragraph in part_2_1:
        add_body(doc, paragraph)
    idef0_path = create_idef0_figure()
    dfd_path = create_dfd_figure()
    use_case_path = create_use_case_figure()
    activity_path = create_activity_figure()
    add_picture_paragraph(doc, idef0_path)
    add_figure_caption(doc, "Рисунок 2.1 - Контекстна IDEF0-модель системи")
    add_picture_paragraph(doc, dfd_path)
    add_figure_caption(doc, "Рисунок 2.2 - DFD-модель обміну даними")
    add_picture_paragraph(doc, use_case_path)
    add_figure_caption(doc, "Рисунок 2.3 - UML use case diagram")
    add_picture_paragraph(doc, activity_path)
    add_figure_caption(doc, "Рисунок 2.4 - UML activity diagram")
    add_table_caption(doc, "Таблиця 2.1 - Перелік моделей та їх призначення")
    add_table(
        doc,
        ["Модель", "Призначення"],
        [
            ["IDEF0", "Опис функціональних блоків і зовнішніх впливів на систему."],
            ["DFD", "Відображення потоків даних між користувачами, сервісами та БД."],
            ["UML use case", "Фіксація ролей і сценаріїв взаємодії з системою."],
            ["UML activity", "Покроковий опис логіки окремих сценаріїв."],
            ["UML class", "Формалізація сутностей і зв’язків доменної моделі."],
            ["UML sequence", "Показ обміну повідомленнями між компонентами."],
        ],
        widths=[4, 12],
    )

    add_heading_2(doc, "2.2 ER-модель бази даних і нормалізація до 3 нормальної форми")
    part_2_2 = [
        "Логічна модель даних побудована за реляційним принципом і охоплює сутності, необхідні для підтримки всіх ключових процесів гуртожитку. Центральними таблицями є Roles, Users і Rooms, які задають структуру доступу та просторову організацію проживання. Від них залежать таблиці ремонтних заявок, перепусток, фінансових нарахувань, платежів, дисциплінарних записів, сесій, логів і файлів.",
        "Нормалізація до третьої нормальної форми зменшує дублювання, усуває аномалії оновлення і спрощує підтримку цілісності даних. Довідникові значення винесено в окремі таблиці, а транзакційні події збережено як окремі сутності з посиланнями на користувачів і пов’язані об’єкти. Для довготривалого зберігання історії використано soft delete та аудиторські журнали замість фізичного видалення критичних даних.",
        "Зв’язок між сутностями реалізовано через зовнішні ключі. Роль має багато користувачів, кімната може містити кількох мешканців, студент може мати кілька нарахувань і кілька платежів, перепустка може мати журнал входу і виходу, а ремонтна заявка пов’язується з категорією, кімнатою та виконавцем. Така структура дозволяє отримувати повну картину стану гуртожитку без зайвих дублювань.",
    ]
    for paragraph in part_2_2:
        add_body(doc, paragraph)
    er_path = create_er_figure()
    add_picture_paragraph(doc, er_path)
    add_figure_caption(doc, "Рисунок 2.5 - ER-модель бази даних")
    add_table_caption(doc, "Таблиця 2.2 - Основні сутності логічної моделі")
    add_table(
        doc,
        ["Сутність", "Призначення"],
        [
            ["Roles", "Довідник ролей користувачів."],
            ["Users", "Облікові записи мешканців і персоналу."],
            ["Rooms", "Кімнатний фонд гуртожитку."],
            ["RepairTickets", "Заявки на ремонт і технічне обслуговування."],
            ["GuestPasses", "Гостьові перепустки з кодом доступу."],
            ["PassLogs", "Журнал фактичних входів і виходів гостей."],
            ["StudentCharges", "Нарахування за проживання."],
            ["Payments", "Реєстрація і підтвердження оплат."],
            ["RelocationRequests", "Заяви на переселення."],
            ["Violations", "Журнал дисциплінарних зауважень."],
            ["Files", "Метадані прикріплених файлів."],
            ["UserSessions", "Активні сесії користувачів."],
            ["RefreshTokens", "Оновлювальні токени для безпечної авторизації."],
            ["AuditLogs", "Службовий журнал дій."],
            ["SecurityEvents", "Журнал подій безпеки."],
        ],
        widths=[5, 11],
    )

    add_heading_2(doc, "2.3 Макет користувацького інтерфейсу")
    part_2_3 = [
        "Користувацький інтерфейс побудовано за принципом рольових кабінетів. Після входу користувач потрапляє до загального каркаса системи, де відображаються навігація, поточний розділ, повідомлення і службові показники. Така організація дозволяє не перевантажувати стартову сторінку зайвими елементами і одночасно швидко переключати користувача між функціональними модулями.",
        "Візуальна мова інтерфейсу базується на картках, м’яких контрастах і адаптивному компонуванні. Це робить систему зручною як на настільному комп’ютері, так і на мобільному пристрої. Для основних процесів використовуються таблиці, панелі деталей, модальні вікна, висувні панелі та форми вводу. Така структура відповідає характеру даних: короткі операції виконуються у формі, а аналітична інформація відображається у вигляді карток і таблиць.",
        "Для студента макет містить огляд проживання, сторінки заявок, перепусток і фінансів. Для коменданта передбачено екран заселеності, переселень і дисциплінарних зауважень. Майстер працює з дошкою ремонтних заявок, охоронець - зі сканером перепусток, бухгалтер - з платежами, а адміністратор - із користувачами та довідниками. Таким чином, кожна роль отримує окремий компактний набір екранів.",
        "Під час проєктування інтерфейсу враховано потребу в чітких підказках, зрозумілих станах кнопок, попередженні про помилки введення та адаптивності для різних розмірів екрана. У результаті інтерфейс не тільки відображає дані, а й допомагає користувачу виконувати дію без зайвих кроків.",
    ]
    for paragraph in part_2_3:
        add_body(doc, paragraph)

    add_heading_2(doc, "2.4 Структура сайту")
    part_2_4 = [
        "Структура сайту побудована за маршрутовим принципом і включає сторінку входу та окремий набір маршрутів для кожної ролі. Після авторизації користувач працює в межах /app, а навігація динамічно підлаштовується під його права доступу. Такий підхід зменшує кількість непотрібних елементів у меню та підтримує логіку розмежування доступу.",
        "Основні маршрути для студента: оглядова сторінка, ремонтні заявки, гостьові перепустки і фінанси. Для коменданта: заселеність, переселення, дисципліна. Для майстра: список ремонтів і деталі заявки. Для охорони: термінал перевірки гостя. Для бухгалтера: платежі. Для адміністратора: студенти, персонал і довідники. Важливо, що кожен модуль підвантажується окремо, тому стартовий набір сторінок залишається компактним.",
        "Така структура добре узгоджується з принципом модульного моноліту на сервері та з роллю окремих UI-секцій на клієнті. Вона забезпечує зрозумілу навігацію, полегшує тестування окремих сценаріїв і створює умови для подальшого розширення системи.",
    ]
    for paragraph in part_2_4:
        add_body(doc, paragraph)

    add_heading_2(doc, "2.5 Алгоритми, що використовувались")
    part_2_5 = [
        "У системі використано не стільки складні математичні моделі, скільки формалізовані алгоритми бізнес-логіки. До них належать перевірка прав доступу за роллю, контроль дійсності сесії, оновлення статусу ремонтної заявки, перевірка наявності місць у кімнаті, валідація гостьової перепустки та підтвердження платежу. Саме ці правила визначають коректну поведінку системи в реальних сценаріях.",
        "Алгоритм авторизації передбачає перевірку облікових даних, створення сесії, видачу access- і refresh-токенів, а також фіксацію CSRF-токена. Алгоритм роботи з сесією враховує період неактивності, що обмежує час життя облікового запису у браузері. Алгоритм перевірки гостей базується на унікальному коді перепустки і часовому інтервалі її дії, а алгоритм підтвердження платежу - на зміні статусу операції та оновленні балансу студента.",
        "Окремо слід згадати допоміжне використання генеративних моделей на етапі підготовки текстових матеріалів і уточнення сценаріїв інтерфейсу. Таке використання обмежувалося допоміжними промптами для формування варіантів формулювань, підказок і тестових кейсів, після чого всі результати вручну перевірялися та адаптувалися до реальної предметної області. Отже, штучний інтелект виступав як інструмент прискорення підготовчої роботи, а не як джерело автономних рішень у системі.",
    ]
    for paragraph in part_2_5:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading_1(doc, "РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ")
    add_heading_2(doc, "3.1 Обґрунтування вибору стеку розробки")
    part_3_1 = [
        "Для клієнтської частини обрано React з TypeScript, Vite, React Router, TanStack Query, React Hook Form і Zod. Такий набір засобів дає змогу створити швидкий SPA-інтерфейс, розділити роботу з даними і компонуванням, реалізувати lazy loading для окремих сторінок і забезпечити строгішу перевірку форм та типів.",
        "Серверна частина побудована на .NET 8 та ASP.NET Core з використанням Entity Framework Core і PostgreSQL. Це рішення добре підходить для системи з рольовим доступом, транзакційними операціями, журналюванням подій і великою кількістю пов’язаних сутностей. PostgreSQL забезпечує надійне збереження даних, а EF Core спрощує роботу з реляційною моделлю на рівні доменних класів.",
        "Для автентифікації використано JWT-обмін через cookie, механізм refresh token, CSRF-захист та rate limiting для авторизаційних запитів. Для зберігання файлів передбачено окремий сервіс сховища, а для розгортання - Docker Compose. Такий стек створює збалансоване поєднання швидкості розробки, безпеки і простоти супроводу.",
    ]
    for paragraph in part_3_1:
        add_body(doc, paragraph)

    add_heading_2(doc, "3.2 Реалізація бази даних")
    part_3_2 = [
        "Фізична модель бази даних відповідає логічній схемі, сформованій на етапі проєктування. Таблиці мають первинні ключі типу UUID, а більшість зв’язків реалізовано через зовнішні ключі з явно визначеною поведінкою при видаленні: Restrict, SetNull або каскадне оновлення там, де це безпечно з точки зору предметної області. Для ключових полів визначено унікальні індекси, щоб запобігати дублюванню email, номерів кімнат, кодів перепусток і хешів refresh token.",
        "У моделі передбачено м’яке видалення через ознаку IsActive, а також службові поля CreatedAt, UpdatedAt, CreatedBy і UpdatedBy. Такий підхід дозволяє зберігати історію змін і не втрачати контекст при аудиті. Окремі сутності, наприклад AuditLogs і SecurityEvents, ведуть журнал критичних дій і спрощують розслідування інцидентів.",
        "Нижче подано скорочений опис основних таблиць фізичної моделі. Деталізовані специфікації полів, індексів і зв’язків доцільно винести до додатків або в технічну документацію проєкту.",
    ]
    for paragraph in part_3_2:
        add_body(doc, paragraph)
    add_table_caption(doc, "Таблиця 3.1 - Основні таблиці фізичної моделі")
    add_table(
        doc,
        ["Таблиця", "Ключові поля", "Призначення"],
        [
            ["Users", "Id, RoleId, RoomId, FullName, Email, Phone, PasswordHash, Balance", "Облікові записи всіх ролей."],
            ["Rooms", "Id, RoomNumber, Floor, Capacity, MonthlyRate, IsUnderRepair", "Кімнатний фонд гуртожитку."],
            ["RepairTickets", "CategoryId, CreatedByUserId, AssignedToUserId, RoomId, Title, Status, Priority", "Заявки на ремонт."],
            ["GuestPasses", "HostUserId, GuestFullName, GuestDocument, ValidFrom, ValidTo, AccessCode", "Гостьові перепустки."],
            ["PassLogs", "PassId, GuardUserId, EntryTime, ExitTime, Remarks", "Журнал проходу гостей."],
            ["StudentCharges", "UserId, Title, Amount, PaidAmount, DueDate, IsSettled", "Нарахування за проживання."],
            ["Payments", "UserId, ChargeId, Amount, PaymentMethod, TransactionStatus, PaidAt", "Облік оплат."],
            ["UserSessions", "UserId, CsrfToken, ExpiresAt, LastActivityAt, IsRevoked", "Сесії користувачів."],
            ["RefreshTokens", "SessionId, TokenHash, ExpiresAt, RevokedAt, JwtId", "Оновлювальні токени."],
            ["AuditLogs", "EntityName, EntityId, Action, UserId, PayloadJson", "Журнал аудиту."],
            ["SecurityEvents", "UserId, EventType, Details, IpAddress, UserAgent", "Події безпеки."],
        ],
        widths=[4, 7, 5],
    )

    add_heading_2(doc, "3.3 Дизайн користувацького інтерфейсу")
    part_3_3 = [
        "Дизайн інтерфейсу реалізовано в сучасній картковій стилістиці з м’якими градієнтами, адаптивними сітками та виразною типографікою. Основними шрифтами стали Manrope для тексту і Fraunces для акцентних заголовків. Така комбінація формує стриманий, але виразний вигляд, що добре підходить для адміністративної системи.",
        "Головна сторінка кабінету показує ім’я користувача, його роль, контактні дані, прив’язку до кімнати та індикатори активних модулів. Для кожної ролі передбачено окремі секції: студентська, комендантська, майстерня, охоронна, бухгалтерська та адміністративна. Це дозволяє зберігати єдину візуальну мову, не змішуючи між собою різні сценарії.",
        "Форми вводу, таблиці, модальні діалоги і висувні панелі допомагають подати складні набори даних у зручнішій формі. Для критичних дій передбачено підтвердження або окреме вікно редагування. Для мобільних пристроїв використано адаптивні навігаційні елементи і компактне компонування, щоб зберегти читабельність на невеликих екранах.",
    ]
    for paragraph in part_3_3:
        add_body(doc, paragraph)

    add_heading_2(doc, "3.4 Розроблення заходів для захисту інформації")
    part_3_4 = [
        "Захист інформації побудовано на поєднанні технічних і організаційних заходів. Паролі користувачів хешуються алгоритмом Argon2id, refresh token не зберігаються у відкритому вигляді, а сесії мають обмежений час бездіяльності. Access token передається через cookie, а для небезпечних методів запитів застосовується CSRF-захист через окремий токен і перевірку заголовка.",
        "Рольовий доступ реалізовано на рівні політик ASP.NET Core, що дозволяє обмежувати доступ до окремих контролерів і операцій. Окрім того, перевірка валідності сесії включає контроль token version, що дає змогу примусово завершувати сесії після змін облікового запису чи видалення користувача. Для авторизаційних запитів застосовано rate limiting, щоб зменшити ризик підбору пароля.",
        "Додатково ведуться журнали аудиту і подій безпеки, зберігаються IP-адреси та user-agent, а при роботі з файлами перевіряється належність користувача до активної сесії. Такі заходи доречно описувати в дипломній роботі як мінімальний набір засобів захисту для системи, яка працює з персональними та фінансовими даними.",
    ]
    for paragraph in part_3_4:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading_1(doc, "РОЗДІЛ 4 ТЕСТУВАННЯ ТА ВПРОВАДЖЕННЯ")
    add_heading_2(doc, "4.1 Обґрунтування вибору методів тестування")
    part_4_1 = [
        "Для системи обрано поєднання модульного, інтеграційного та сценарного тестування. Модульне тестування доцільне для перевірки окремих алгоритмів і правил, наприклад контролю доступу, валідації перепусток або обробки платежів. Інтеграційне тестування потрібне для перевірки взаємодії клієнтської частини, API і бази даних.",
        "Сценарне тестування дозволяє перевірити реальні бізнес-кейси: вхід студента, створення заявки, погодження переселення, роботу сканера на прохідній і підтвердження платежу бухгалтером. Оскільки система багаторольова, особливо важливо перевірити доступність лише тих функцій, які відповідають конкретній ролі. Це зменшує ризик помилок авторизації і логічних суперечностей у поведінці інтерфейсу.",
    ]
    for paragraph in part_4_1:
        add_body(doc, paragraph)

    add_heading_2(doc, "4.2 Формування тест-плану")
    part_4_2 = [
        "Тест-план повинен охоплювати як позитивні, так і негативні сценарії. Для кожної ролі доцільно перевірити вхід у систему, виконання основних дій, реакцію на помилкові дані, збереження результату в БД та коректність повідомлень в інтерфейсі. Тест-кейси і чек-листи зручно винести до додатків, а в основному тексті подати лише стислий огляд найбільш важливих сценаріїв.",
    ]
    for paragraph in part_4_2:
        add_body(doc, paragraph)
    add_table_caption(doc, "Таблиця 4.1 - Приклад фрагмента тест-плану")
    add_table(
        doc,
        ["ID", "Сценарій", "Очікуваний результат"],
        [
            ["TC-01", "Студент входить у систему з коректними обліковими даними", "Відкривається особистий кабінет студента."],
            ["TC-02", "Студент створює заявку на ремонт", "Заявка з’являється у списку, статус - Нова."],
            ["TC-03", "Комендант погоджує заявку на переселення", "Статус змінюється на Схвалено, кімната оновлюється."],
            ["TC-04", "Охоронець перевіряє чинну перепустку", "Система підтверджує право на прохід і фіксує подію."],
            ["TC-05", "Бухгалтер підтверджує платіж", "Змінюється статус оплати і зменшується заборгованість."],
            ["TC-06", "Користувач вводить неправильний пароль", "Система відхиляє вхід і реєструє безпекову подію."],
        ],
        widths=[2, 8, 6],
    )

    add_heading_2(doc, "4.3 Впровадження системи")
    part_4_3 = [
        "Для впровадження системи передбачено контейнерний сценарій, у якому окремо запускаються база даних PostgreSQL, серверна частина, клієнтська частина і допоміжні сервіси. У складі `docker-compose.yml` визначено контейнери для `postgres`, `minio`, `api` і `web`. Такий підхід спрощує локальне розгортання, відтворюваність середовища і демонстрацію системи на захисті.",
        "Серверна частина збирається з Dockerfile на базі .NET 8, а вебклієнт - з Node.js у двох режимах: розробки та збірки. У конфігурації вказано змінні для підключення до бази даних, авторизації та файлового сховища. Це дозволяє відокремити код застосунку від середовища запуску і полегшує подальше перенесення системи на інший сервер.",
        "Під час впровадження важливо перевірити, чи коректно створюються початкові довідники, чи доступні тестові облікові записи ролей і чи зберігаються файли та журнали подій. Саме ця частина демонструє практичну готовність системи до роботи в реальному середовищі гуртожитку.",
    ]
    for paragraph in part_4_3:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading_1(doc, "ВИСНОВКИ")
    conclusions = [
        "У ході виконання кваліфікаційної роботи проаналізовано предметну область студентського гуртожитку та визначено ключові процеси, що потребують цифрової підтримки. Це дозволило сформувати цілісне бачення системи, у якій проживання, технічне обслуговування, контроль доступу і фінансовий облік поєднані в одному інформаційному просторі.",
        "На етапі проєктування побудовано моделі IDEF0, DFD та UML, сформовано ER-модель бази даних, виконано нормалізацію даних до третьої нормальної форми та визначено структуру основних сутностей. Окремо розроблено концепцію інтерфейсу і структуру сайту з урахуванням рольової моделі доступу.",
        "У програмній частині обґрунтовано вибір стеку розробки на базі React, ASP.NET Core, PostgreSQL і Docker, а також описано механізми захисту інформації: Argon2id, JWT, CSRF, контроль сесій і журналювання подій. Це забезпечує належний рівень безпеки та підтримуваності системи.",
        "Загалом отримана модель інформаційної системи придатна для подальшої реалізації та практичного впровадження у студентському гуртожитку. Вона може стати основою для розширення функціональності, інтеграції з додатковими сервісами і подальшої автоматизації адміністративних процесів.",
    ]
    for paragraph in conclusions:
        add_body(doc, paragraph)

    doc.add_page_break()
    add_heading_1(doc, "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ")
    references = [
        "1. Методичні вказівки до виконання кваліфікаційної роботи бакалавра для здобувачів вищої освіти спеціальності 122 «Комп’ютерні науки». Полтава: НУПП, 2024. 55 с.",
        "2. Постанова Кабінету Міністрів України «Про затвердження Примірного положення про користування гуртожитками» № 498 від 20 червня 2018 року. URL: https://zakon.rada.gov.ua/go/498-2018-%D0%BF",
        "3. Положення про особливості користування гуртожитками закладів фахової передвищої та вищої освіти. URL: https://zakon.rada.gov.ua/go/z0114-20",
        "4. Міністерство освіти і науки України. Реєстрація місця проживання в гуртожитках ЗВО вже можлива в Дії. URL: https://mon.gov.ua/news/reiestratsiia-mistsia-prozhyvannia-v-hurtozhytkakh-zvo-vzhe-mozhlyva-v-dii",
        "5. StarRez. Student Housing Management Software. URL: https://www.starrez.com/",
        "6. Entrata. Student Housing Management Software. URL: https://www.entrata.com/student.",
        "7. RMS Cloud. Student accommodation management software. URL: https://www.rmscloud.com/solutions/student-accomodation-management-software",
        "8. Roompact. URL: https://www.roompact.com/",
        "9. Microsoft Learn. ASP.NET Core documentation. URL: https://learn.microsoft.com/en-us/aspnet/core/",
        "10. Microsoft Learn. Entity Framework Core documentation. URL: https://learn.microsoft.com/en-us/ef/core/",
        "11. React. Official documentation. URL: https://react.dev/",
        "12. React Router. Documentation. URL: https://reactrouter.com/",
        "13. TanStack Query. Documentation. URL: https://tanstack.com/query/latest",
        "14. PostgreSQL Global Development Group. PostgreSQL Documentation. URL: https://www.postgresql.org/docs/",
        "15. Docker Documentation. URL: https://docs.docker.com/",
        "16. Vite. Guide and API Reference. URL: https://vite.dev/",
        "17. OWASP Foundation. Cheat Sheet Series. URL: https://cheatsheetseries.owasp.org/",
    ]
    for ref in references:
        add_reference(doc, ref)

    doc.add_page_break()
    add_heading_1(doc, "ДОДАТКИ")
    appendices = [
        "Додаток А - Перелік ілюстрацій до моделі системи: контекстна IDEF0-діаграма, DFD-рівні, UML use case, activity, class і sequence diagrams.",
        "Додаток Б - ER-модель бази даних та таблиці нормалізації до 3 нормальної форми.",
        "Додаток В - Фрагменти інтерфейсу користувача і структура основних сторінок вебсайту.",
        "Додаток Г - Тест-кейси та чек-листи для основних сценаріїв роботи системи.",
        "Додаток Д - Фрагменти вихідного коду ключових модулів: автентифікація, база даних, сервіси бізнес-логіки та контейнерне розгортання.",
    ]
    for item in appendices:
        add_body(doc, item)

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT_FILE)
    print(f"Saved to {OUT_FILE}")


if __name__ == "__main__":
    main()
