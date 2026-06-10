from __future__ import annotations

from pathlib import Path
import gc
import shutil
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.build_diploma_report import (  # noqa: E402
    add_canvas_title,
    create_canvas,
    create_class_figure,
    create_er_figure,
    create_sequence_figure,
    create_wireframe_figure,
    draw_arrow,
    draw_box,
    load_font,
    save_image,
)


DOCX_PATH = ROOT / "labs_reports" / "Diploma_Report_Student_Dormitory.docx"
OUTPUT_COPY = ROOT / "output" / "doc" / "Diploma_Report_Student_Dormitory.docx"


def find_paragraph(document: Document, prefix: str, occurrence: int = 1):
    count = 0
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            count += 1
            if count == occurrence:
                return paragraph
    raise ValueError(f"Paragraph starting with {prefix!r} not found")


def insert_before(paragraph, text: str | None = None):
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_element
    if text is not None:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_after(paragraph, text: str | None = None):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_element
    if text is not None:
        new_paragraph.add_run(text)
    return new_paragraph


def delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def style_heading1(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def style_heading2(paragraph):
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def style_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def style_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_picture_before(paragraph, image_path: Path, width_cm: float = 15.5):
    picture_paragraph = insert_before(paragraph)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return picture_paragraph


def add_picture_after(paragraph, image_path: Path, width_cm: float = 15.5):
    picture_paragraph = insert_after(paragraph)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return picture_paragraph


def build_extra_paragraphs(anchor, texts: list[str]):
    last = anchor
    for text in texts:
        last = insert_after(last, text)
        style_body(last)
    return last


def create_site_map_figure() -> Path:
    image, draw = create_canvas()
    add_canvas_title(draw, "Схема маршрутизації та навігації застосунку", "Рольова структура переходів у межах /app")

    title_font = load_font(22, bold=True)
    box_font = load_font(18, bold=True)
    small_font = load_font(15)

    draw.rounded_rectangle((860, 150, 1340, 260), radius=24, fill=(239, 246, 255), outline=(59, 130, 246), width=4)
    draw.text((995, 185), "Login", font=title_font, fill=(30, 41, 59))
    draw.rounded_rectangle((860, 320, 1340, 430), radius=24, fill=(236, 253, 245), outline=(16, 185, 129), width=4)
    draw.text((960, 355), "/app shell", font=title_font, fill=(30, 41, 59))
    draw_arrow(draw, (1100, 260), (1100, 320), color=(59, 130, 246), width=5)

    role_boxes = [
        ("Студент", (120, 560, 470, 670), (168, 85, 247)),
        ("Комендант", (580, 560, 930, 670), (245, 158, 11)),
        ("Майстер", (1040, 560, 1390, 670), (14, 165, 233)),
        ("Охорона", (1500, 560, 1850, 670), (34, 197, 94)),
        ("Бухгалтер", (350, 900, 700, 1010), (220, 38, 38)),
        ("Адміністратор", (930, 900, 1280, 1010), (99, 102, 241)),
    ]

    for role, box, outline in role_boxes:
        draw_box(draw, box, role, box_font, fill=(248, 250, 252), outline=outline, width=4)

    shared_routes = [
        ("Огляд", (120, 760, 260, 830)),
        ("Заявки", (285, 760, 425, 830)),
        ("Профіль", (450, 760, 590, 830)),
    ]
    for text, box in shared_routes:
        draw_box(draw, box, text, small_font, fill=(255, 255, 255), outline=(148, 163, 184), width=2, radius=16)

    role_routes = {
        "Студент": ["Кімната", "Перепустки", "Фінанси"],
        "Комендант": ["Заселеність", "Переселення", "Дисципліна"],
        "Майстер": ["Ремонтні заявки", "Черга робіт"],
        "Охорона": ["Перевірка пропуску", "Журнал входу"],
        "Бухгалтер": ["Нарахування", "Платежі"],
        "Адміністратор": ["Користувачі", "Довідники", "Ролі"],
    }

    for role, box, outline in role_boxes:
        x0, y0, x1, y1 = box
        label_y = y1 + 18
        for index, item in enumerate(role_routes[role]):
            item_box = (x0 + 20 + index * 110, label_y, x0 + 110 + index * 110, label_y + 44)
            draw_box(draw, item_box, item, small_font, fill=(255, 255, 255), outline=outline, width=2, radius=12)
        draw_arrow(draw, (1100, 430), ((x0 + x1) // 2, y0), color=outline, width=4)

    draw.text((145, 1110), "Після авторизації користувач отримує лише ті маршрути, які відповідають його ролі.", font=small_font, fill=(51, 65, 85))
    draw.text((145, 1150), "Спільний каркас сторінки зберігає єдину навігацію, а внутрішні розділи змінюються без перезавантаження застосунку.", font=small_font, fill=(51, 65, 85))

    return save_image(image, "figure_2_9_sitemap.png")


def main() -> None:
    class_path = create_class_figure()
    sequence_path = create_sequence_figure()
    wireframe_path = create_wireframe_figure()
    er_path = create_er_figure()
    sitemap_path = create_site_map_figure()

    doc = Document(DOCX_PATH)

    # Table of contents updates.
    for prefix, text in [
        ("2.1 Проєктування функціоналу та структури системи керування гуртожитком", "2.1 Проєктування функціоналу та структури системи керування гуртожитком      15"),
        ("2.2 Проєктування бази даних", "2.2 Проєктування бази даних      20"),
        ("2.3 Макети користувацького інтерфейсу", "2.3 Макети користувацького інтерфейсу      24"),
        ("2.4 Структура сайту/застосунку", "2.4 Структура сайту/застосунку      28"),
        ("РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ", "РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ      31"),
        ("3.1 Обґрунтування вибору стеку розробки", "3.1 Обґрунтування вибору стеку розробки      31"),
        ("3.2 Реалізація бази даних", "3.2 Реалізація бази даних      33"),
        ("3.3 Реалізація основної функціональності", "3.3 Реалізація основної функціональності      35"),
        ("3.4 Дизайн користувацького інтерфейсу", "3.4 Дизайн користувацького інтерфейсу      37"),
        ("3.5 Розроблення заходів для захисту інформації", "3.5 Розроблення заходів для захисту інформації      39"),
        ("РОЗДІЛ 4 ТЕСТУВАННЯ ТА ВПРОВАДЖЕННЯ", "РОЗДІЛ 4 ТЕСТУВАННЯ ТА ВПРОВАДЖЕННЯ      41"),
        ("4.1 Обґрунтування вибору методів тестування", "4.1 Обґрунтування вибору методів тестування      41"),
        ("4.2 Формування тест-плану", "4.2 Формування тест-плану      42"),
        ("4.3 Тест-кейси та чек-листи", "4.3 Тест-кейси та чек-листи      44"),
        ("4.4 Впровадження системи", "4.4 Впровадження системи      46"),
        ("ВИСНОВКИ", "ВИСНОВКИ      48"),
        ("СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ      49"),
        ("ДОДАТКИ", "ДОДАТКИ      51"),
    ]:
        try:
            paragraph = find_paragraph(doc, prefix, 1)
        except ValueError:
            continue
        paragraph.text = text
        if prefix.startswith(("РОЗДІЛ", "ВИСНОВКИ", "СПИСОК", "ДОДАТКИ")):
            style_heading1(paragraph)
        else:
            style_body(paragraph)

    # Chapter 2 heading updates.
    chapter_2_1 = find_paragraph(doc, "2.1 Проєктування функціоналу та структури системи керування гуртожитком", 2)
    chapter_2_2 = find_paragraph(doc, "2.2 ER-модель бази даних і нормалізація до 3 нормальної форми", 1)
    chapter_2_3 = find_paragraph(doc, "2.3 Макет користувацького інтерфейсу", 1)
    chapter_2_4 = find_paragraph(doc, "2.4 Структура сайту", 1)
    obsolete_tail_paragraphs = [doc.paragraphs[index] for index in range(198, 210)]
    chapter_3 = find_paragraph(doc, "РОЗДІЛ 3 ПРОГРАМНА РЕАЛІЗАЦІЯ", 1)
    obsolete_algorithm_paragraphs = [doc.paragraphs[index] for index in (195, 196, 197)]

    chapter_2_2.text = "2.2 Проєктування бази даних"
    chapter_2_3.text = "2.3 Макети користувацького інтерфейсу"
    chapter_2_4.text = "2.4 Структура сайту/застосунку"
    style_heading2(chapter_2_2)
    style_heading2(chapter_2_3)
    style_heading2(chapter_2_4)

    # Expand subsection 2.1.
    chapter_2_1_texts = [
        "Проєктування системи керування гуртожитком починається з чіткого визначення меж предметної області та переліку процесів, які повинні бути підтримані цифровим середовищем. Для такого проєкту важливо одночасно врахувати інтереси студентів, коменданта, майстра, охорони, бухгалтера й адміністратора, оскільки кожна роль працює з власним набором дій, але використовує спільні дані про кімнати, заявки, перепустки, платежі та користувачів.",
        "Функціональна модель дає змогу не лише перелічити можливості системи, а й показати, як саме окремі процеси взаємопов’язані між собою. У межах гуртожитку це означає, що авторизація, перегляд заселення, подання заявок, погодження переселення, перевірка перепусток і фінансові операції розглядаються як частини єдиного потоку управління, а не як ізольовані сценарії.",
        "Контекстна IDEF0-модель використовується для фіксації функції верхнього рівня та її оточення. Входами для системи є дані про користувачів, кімнати, звернення, перепустки й платежі, керуючими впливами - правила проживання, рольова модель, часові обмеження та політики безпеки, механізмами - вебзастосунок, серверна логіка і база даних, а виходами - оновлена інформація, повідомлення та звіти для персоналу.",
        "Декомпозиція IDEF0 доцільна тому, що дозволяє поступово розкрити кожен великий процес на менші функціональні блоки. Для системи гуртожитку такими блоками є управління користувачами, облік кімнат, оформлення ремонтних звернень, обробка гостьових перепусток, облік платежів і ведення службових журналів. Саме така структура допомагає уникнути надмірної складності на етапі реалізації.",
        "DFD-модель зосереджується на русі даних між інтерфейсом, серверною частиною, базою даних та зовнішніми сервісами. Вона показує, як запит від користувача проходить через форму, валідацію, сервісну логіку та сховище даних, а потім повертається у вигляді таблиці, повідомлення або статусу виконання. Для студентського гуртожитку це особливо важливо, оскільки більшість операцій має чітко визначений життєвий цикл.",
        "На рівні потоків даних окрему увагу приділено сценаріям, у яких результат залежить від кількох перевірок. Наприклад, під час створення заявки на переселення система перевіряє наявність вільних місць, коректність заповнення форми та доступність ролі для виконання дії. Подібним чином працює і перевірка гостьової перепустки, де ключовими є часові межі, унікальний код та належність перепустки конкретному студенту.",
        "UML-діаграми застосовуються для деталізації поведінки системи на рівні користувацьких сценаріїв та об’єктів. Діаграма варіантів використання показує, які функції доступні кожній ролі, діаграма діяльності - послідовність кроків у межах окремого сценарію, діаграма класів - набір сутностей та зв’язків доменної моделі, а діаграма послідовності - порядок обміну повідомленнями між клієнтом, API і базою даних.",
        "Такий набір моделей не дублює інформацію, а доповнює її з різних точок зору. Якщо IDEF0 відповідає на питання, що робить система, то DFD пояснює, як рухаються дані, use case описує, хто і з якою метою використовує функції, activity деталізує кроки виконання, class формалізує структуру даних, а sequence демонструє технічну взаємодію компонентів.",
        "Особливо корисними є моделі для сценаріїв, де відбувається зміна стану об’єкта. Саме так працюють ремонтні заявки, платежі, перепустки та переселення: кожен з цих процесів має початковий стан, проміжні перевірки, відповідальну роль і очікуваний результат. Візуальне моделювання допомагає зафіксувати ці переходи ще до написання коду.",
        "Окремі діаграми також допомагають узгодити проєктну документацію з майбутньою реалізацією фронтенду. Коли на схемі видно, що користувач спершу проходить авторизацію, а далі потрапляє у рольовий кабінет, це безпосередньо впливає на маршрутизацію, побудову меню, перевірку доступу та організацію спільного каркаса застосунку.",
        "Комплект моделей у цьому підрозділі є основою для подальших рішень щодо бази даних, інтерфейсу та структури застосунку. Саме на цьому етапі формується цілісна картина того, як цифрова система повинна підтримувати реальні процеси студентського гуртожитку без втрати логіки та цілісності даних.",
    ]
    chapter_2_1_body = [doc.paragraphs[index] for index in range(154, 166)]
    for paragraph, text in zip(chapter_2_1_body, chapter_2_1_texts):
        paragraph.text = text
        style_body(paragraph)
    chapter_2_1_last = chapter_2_1_body[-1]
    chapter_2_1_last = insert_after(
        chapter_2_1_last,
        "Окреме значення має фіксація переходів між станами об’єктів, оскільки саме зміна стану заявки, перепустки або платежу формує основу для контролю й подальшого аудиту. Завдяки цьому проєктні діаграми відображають не лише статичну структуру, а й динаміку роботи системи у реальних сценаріях.",
    )
    style_body(chapter_2_1_last)

    # Insert class and sequence diagrams before the next subsection.
    class_picture = add_picture_before(chapter_2_2, class_path)
    class_caption = insert_after(class_picture, "Рисунок 2.5 - Діаграма класів предметної області")
    style_caption(class_caption)

    sequence_picture = insert_after(class_caption)
    sequence_picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sequence_run = sequence_picture.add_run()
    sequence_run.add_picture(str(sequence_path), width=Cm(15.5))
    sequence_caption = insert_after(sequence_picture, "Рисунок 2.6 - Діаграма послідовності взаємодії компонентів")
    style_caption(sequence_caption)

    # Expand subsection 2.2.
    chapter_2_2_texts = [
        "Проєктування бази даних у цій роботі спирається на логіку предметної області, тому реляційна схема формується не від набору таблиць, а від процесів, які ця схема повинна підтримувати. Логічна модель охоплює довідники ролей, облікові записи користувачів, кімнати, ремонтні звернення, гостьові перепустки, платежі, службові журнали та інші сутності, що відображають реальну роботу гуртожитку.",
        "Під час побудови логічної моделі важливо відокремити довідникові дані від транзакційних подій. Ролі, типи звернень, категорії кімнат і подібні довідники змінюються рідко, тоді як заявки, платежі, перепустки та записи журналу створюються щодня. Такий поділ спрощує підтримку цілісності та зменшує кількість дублювань у базі.",
        "Нормалізація до третьої нормальної форми дає змогу прибрати надлишкові залежності й уникнути аномалій вставки, оновлення та видалення. Для проєкту гуртожитку це означає, що інформація про мешканця, його кімнату, фінансовий стан і службові звернення зберігається в окремих, логічно пов’язаних таблицях, а не повторюється в кожному записі.",
    ]
    chapter_2_2_body = [doc.paragraphs[index] for index in range(176, 179)]
    for paragraph, text in zip(chapter_2_2_body, chapter_2_2_texts):
        paragraph.text = text
        style_body(paragraph)
    chapter_2_2_last = insert_after(
        chapter_2_2_body[-1],
        "Зв’язки між сутностями реалізуються через зовнішні ключі з чітко визначеною поведінкою при зміні пов’язаних записів. Наприклад, роль має багато користувачів, кімната може бути пов’язана з кількома мешканцями, студент може мати кілька заявок і кілька платежів, а ремонтне звернення пов’язується з кімнатою та відповідальним виконавцем. Така схема дозволяє вести історію процесів без втрати контексту.",
    )
    style_body(chapter_2_2_last)
    chapter_2_2_last = insert_after(
        chapter_2_2_last,
        "Під час переходу від логічної до фізичної моделі для ключових полів використовуються UUID, унікальні індекси та службові атрибути часу створення і зміни запису. Це особливо важливо для таких значень, як електронна пошта, номер кімнати, код перепустки або хеш оновлювального токена, де дублювання неприпустиме з погляду як функціональності, так і безпеки.",
    )
    style_body(chapter_2_2_last)
    chapter_2_2_last = insert_after(
        chapter_2_2_last,
        "Окремо передбачено м’яке видалення записів, службові ознаки активності та аудиторські журнали. Такий підхід дозволяє зберігати історію змін і відновлювати контекст під час перевірок або розслідування інцидентів, що є важливим для адміністративної системи з персональними і фінансовими даними.",
    )
    style_body(chapter_2_2_last)

    er_picture = insert_after(sequence_caption)
    er_picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er_run = er_picture.add_run()
    er_run.add_picture(str(er_path), width=Cm(15.5))
    er_caption = insert_after(er_picture, "Рисунок 2.7 - ER-модель бази даних")
    style_caption(er_caption)
    table_caption = insert_after(er_caption, "Таблиця 2.2 - Основні сутності логічної моделі")
    style_caption(table_caption)

    # The table itself is already present in the document; we only keep the caption coherent.
    # Expand subsection 2.3.
    chapter_2_3_texts = [
        "Макети користувацького інтерфейсу розробляються у форматі wireframe, тобто як структурні схеми сторінок без фінального візуального стилю. Такий підхід дає змогу зосередитися на ієрархії блоків, розміщенні ключових елементів та логіці переходів, не відволікаючись на декоративні деталі, які не впливають на функціональну зрозумілість інтерфейсу.",
        "Базова композиція сторінки охоплює хедер, навігаційне меню, робочу область, інформаційні панелі та зону дій користувача. У верхній частині доцільно розміщувати назву розділу, коротку інформацію про роль, сповіщення і кнопку виходу, а в основній частині - таблиці, картки, форми та модальні вікна, які підтримують конкретний сценарій роботи.",
        "Для робочого місця студента wireframe містить огляд кімнати, фінансового стану, список заявок і блок швидких дій для подання звернення або перегляду перепусток. Для коменданта основними елементами є заселеність, списки кімнат, заявки на переселення та дисциплінарні записи. Для майстра окремо виділяється дошка ремонтних звернень, а для охорони - проста форма перевірки перепустки.",
        "Wireframe-модель є корисною ще й тому, що дозволяє перевірити, чи не перевантажено екран другорядними елементами. На цьому етапі легко побачити, чи достатньо помітні основні кнопки, чи правильно розміщено навігацію і чи зрозумілий шлях користувача від перегляду інформації до виконання дії. Саме тому макети є повноцінним етапом проєктування, а не допоміжним ескізом.",
    ]
    chapter_2_3_body = [doc.paragraphs[index] for index in range(183, 187)]
    for paragraph, text in zip(chapter_2_3_body, chapter_2_3_texts):
        paragraph.text = text
        style_body(paragraph)
    chapter_2_3_last = chapter_2_3_body[-1]
    chapter_2_3_last = insert_after(
        chapter_2_3_last,
        "Кабінет бухгалтера орієнтований на таблиці нарахувань і платежів, фільтри за статусами та зрозумілі дії підтвердження. Адміністративний кабінет, своєю чергою, містить окремі області для керування студентами, персоналом, ролями та довідниками. Спільним для всіх варіантів є те, що короткі дії виконуються у формі, а аналітичні дані показуються у вигляді таблиць і карток.",
    )
    style_body(chapter_2_3_last)
    chapter_2_3_last = insert_after(
        chapter_2_3_last,
        "Wireframe також допомагає узгодити візуальну ієрархію елементів із пріоритетами конкретної ролі. Наприклад, для студента першорядним є доступ до особистих дій, а для бухгалтера - до таблиць і фільтрів, тому один і той самий каркас сторінки має різне наповнення залежно від сценарію роботи.",
    )
    style_body(chapter_2_3_last)

    wireframe_picture = insert_after(chapter_2_3_last)
    wireframe_picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wireframe_run = wireframe_picture.add_run()
    wireframe_run.add_picture(str(wireframe_path), width=Cm(15.5))
    wireframe_caption = insert_after(wireframe_picture, "Рисунок 2.8 - Wireframe користувацького інтерфейсу")
    style_caption(wireframe_caption)

    # Expand subsection 2.4.
    chapter_2_4_texts = [
        "Структура сайту або застосунку побудована за маршрутовим принципом і починається зі сторінки входу. Після успішної авторизації користувач потрапляє до спільної оболонки системи, у якій зберігаються хедер, меню, зона повідомлень і робочий простір. Це дозволяє не дублювати каркас сторінки на кожному екрані та зосередити увагу на змісті конкретного модуля.",
        "На рівні маршрутизації система використовує окремі гілки для кожної ролі. Студент працює з оглядом, заявками, перепустками та фінансами; комендант - із заселеністю, переселенням і дисципліною; майстер - з ремонтними заявками; охорона - з терміналом перевірки перепусток; бухгалтер - з платежами; адміністратор - з користувачами та довідниками.",
        "Такий поділ спрощує контроль доступу і допомагає підтримувати зрозумілу навігацію. Користувач бачить лише ті розділи, які належать до його робочого сценарію, а отже не витрачає час на зайві переходи. Це особливо важливо для адміністративної системи, де різні ролі працюють у межах спільних даних, але виконують принципово різні операції.",
        "З точки зору фронтенд-архітектури застосунок є не набором розрізнених сторінок, а узгодженою картою маршрутів. Вкладені маршрути, активне підсвічування пунктів меню та однаковий підхід до підключення сторінок роблять інтерфейс передбачуваним і спрощують подальше розширення системи без повної перебудови структури.",
        "Схема маршрутизації та навігації пов’язує функціональну модель, модель даних і макети інтерфейсу в один логічний ланцюг. Якщо на попередніх етапах визначено ролі, сутності та розміщення елементів, то на цьому етапі стає зрозуміло, як саме користувач рухається між екранами і в якому контексті використовує кожен модуль.",
        "Окрему увагу приділено адаптивності навігації. На великих екранах зручно використовувати постійне меню, тоді як на мобільних пристроях воно може бути згорнутим або винесеним у нижню панель. Такий підхід допомагає зберегти читабельність інтерфейсу і не перевантажує невеликі екрани зайвими елементами.",
        "Отже, структура застосунку підтримує не лише навігацію, а й архітектурну дисципліну проєкту. Вона фіксує межі спільного каркаса, визначає доступні маршрути для кожної ролі та забезпечує послідовний перехід від огляду даних до виконання конкретної операції.",
    ]
    chapter_2_4_body = [doc.paragraphs[index] for index in range(188, 195)]
    for paragraph, text in zip(chapter_2_4_body, chapter_2_4_texts):
        paragraph.text = text
        style_body(paragraph)
    chapter_2_4_last = chapter_2_4_body[-1]
    chapter_2_4_last = insert_after(
        chapter_2_4_last,
        "У підсумку маршрутна карта застосунку узгоджується з ролями та бізнес-процесами гуртожитку: користувач проходить короткий шлях від входу до конкретної дії, а система зберігає єдиний підхід до відображення даних, навігації та перевірки доступу.",
    )
    style_body(chapter_2_4_last)

    sitemap_picture = insert_after(chapter_2_4_last)
    sitemap_picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sitemap_run = sitemap_picture.add_run()
    sitemap_run.add_picture(str(sitemap_path), width=Cm(15.5))
    sitemap_caption = insert_after(sitemap_picture, "Рисунок 2.9 - Схема маршрутизації та навігації застосунку")
    style_caption(sitemap_caption)

    # Remove the obsolete algorithm subsection and its explanatory paragraphs.
    for paragraph in obsolete_algorithm_paragraphs:
        delete_paragraph(paragraph)
    for paragraph in obsolete_tail_paragraphs:
        delete_paragraph(paragraph)

    # Keep the old table caption aligned after inserting new figures.
    for paragraph in [chapter_2_1, chapter_2_2, chapter_2_3, chapter_2_4, chapter_3]:
        style_heading2(paragraph)

    # Save through a temporary file to avoid sharing violations on Windows.
    temp_path = DOCX_PATH.with_name("Diploma_Report_Student_Dormitory.updated.docx")
    doc.save(temp_path)
    del doc
    gc.collect()

    shutil.copy2(temp_path, DOCX_PATH)
    shutil.copy2(temp_path, OUTPUT_COPY)
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
